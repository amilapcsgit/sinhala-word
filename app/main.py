# Sinhala Word 2025 — v6.1 - Multi-Format with Printing Support
"""
A lightweight **Word‑2000‑style** editor refreshed with Windows‑11 Fluent UI:

• **Realtime Singlish → Sinhala** (Buffer-based input like Sintip)
• **Suggestion popup** (up to 9 matches, Tab/1-9/Click to accept)
• **Basic spell‑checker** — unknown Sinhala words are underlined red
• **Two classic toolbars** (Standard & Formatting) ‑or‑ hide them and use the menu
• **Dark / Light theme toggle** (View → Toggle Theme) - Enhanced implementation
• **Multi-format support** — Open and save files in TXT, DOCX, and PDF formats and print them
• **Sinhala keyboard** (on-screen) - Optional, can be hidden behind the main window for convenience
• **On-Screen Keyboard** (optional) - Can be toggled via View → Toggle Keyboard
• **Settings dialog** (accessible via File → Preferences)
• **Customizable font settings**
• **User dictionary management** — Add/edit/delete entries directly from the app
• **Automatic line wrapping** — Text wraps automatically at the end of each line
• **Line numbers** — Display line numbers alongside text for easy reference
• **Font size adjustment** — Easily adjust font size using toolbar buttons or shortcut keys
• **Undo/Redo functionality** — Undo and redo changes made to the document easily

Install dependencies:
```bash
pip install PySide6 PySide6-Fluent-Widgets python-docx reportlab pypdf
```

Run:
```bash
python run.py
```

Pack as EXE:
```bash
pyinstaller --noconfirm --onefile --hidden-import=docx --hidden-import=reportlab --hidden-import=pypdf --add-data "sinhalawordmap.json;." --add-data "dictionary;dictionary" run.py
```
"""
import sys
import os
import json
import re
import gzip
import logging

# --- Dynamic dependency bootstrap ---------------------------------
import importlib.util, subprocess, sys, logging
log = logging.getLogger(__name__)
_REQS = {
    "python-docx": "docx",
    "reportlab": "reportlab", # reportlab is still needed for now if _do_save calls it for other types.
    "pypdf": "pypdf",
    "PySide6.QtPrintSupport": "PySide6.QtPrintSupport" # Added QPrinter dependency
}

def _ensure_deps():
    for pip_name, mod_name in _REQS.items():
        if importlib.util.find_spec(mod_name) is None:
            log.warning(f"Missing '{pip_name}', installing…")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])

_ensure_deps()
# ------------------------------------------------------------------

# Allow limited font fallbacks for better compatibility
# We'll handle font fallbacks more carefully in the code
os.environ["QT_ENABLE_FONT_FALLBACKS"] = "1"
os.environ["QT_FONT_NO_SYSTEM_FALLBACKS"] = "0"
from PySide6.QtWidgets import (
    QApplication, QTextEdit, QFileDialog, QToolBar, QWidget, QVBoxLayout,
    QFontComboBox, QComboBox, QMessageBox, QStatusBar, QLabel, 
    QFrame, QInputDialog, QMainWindow, QPushButton, QHBoxLayout, QSizePolicy,
    QStyledItemDelegate, QMenu, QSplitter, QDialog
)
from PySide6.QtPrintSupport import QPrintDialog, QPrinter, QPrintPreviewDialog
from PySide6.QtGui import QFont, QTextCursor, QTextCharFormat, QColor, QAction, QIcon, QFontDatabase
from PySide6.QtCore import Qt, QPoint, QTimer, QEvent, Slot, QSize, QObject
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

# Import our custom modules
from ui.keyboard import SinhalaKeyboard
from ui.theme_manager import ThemeManager
from app import config
from app.transliterator import SinhalaTransliterator
from app.spellchecker import SinhalaSpellChecker
from app.input_handler import SinhalaInputHandler
from ui.suggestion_popup import SuggestionPopup
from ui.settings_dialog import SettingsDialog
from ui.icons import get_toolbar_icon

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("SinhalaWordProcessor")

# ------------------------------------------------------------------
#  Constants & Global Helpers
# ------------------------------------------------------------------
WORD_PATTERN = re.compile(r'\b\w+\b')  # Compiled regex for word counting

# Custom delegate for better rendering of font sizes in dropdown
class FontSizeDelegate(QStyledItemDelegate):
    def __init__(self, parent=None):
        super().__init__(parent)
    
    def paint(self, painter, option, index):
        # Simply use the default painting with centered alignment
        option.displayAlignment = Qt.AlignCenter
        super().paint(painter, option, index)

# Import the FontManager
from ui.font_manager import FontManager

def configure_font_fallbacks():
    """Configure font fallbacks for better compatibility."""
    try:
        # Allow limited font fallbacks for better compatibility
        # This helps with rendering Sinhala characters properly
        os.environ["QT_ENABLE_FONT_FALLBACKS"] = "1"
        os.environ["QT_FONT_NO_SYSTEM_FALLBACKS"] = "0"
        logging.info("Configured font fallbacks for better compatibility")
    except Exception as e:
        logging.error(f"Failed to configure font fallbacks: {e}")

def load_sinhala_fonts():
    """Load Sinhala fonts using the FontManager."""
    # Initialize the font manager (singleton)
    font_manager = FontManager()
    
    # Load fonts
    font_manager.load_fonts()
    
    # Return the list of available fonts
    return font_manager.get_available_fonts()





# Phonetic fallback definitions
VOW = {"aa":"ා","a":"","ae":"ැ","aae":"ෑ","i":"ි","ii":"ී","u":"ු","uu":"ූ",
       "ru":"ෘ","ruu":"ෲ","e":"ෙ","ee":"ේ","o":"ො","oo":"ෝ","au":"ෞ","":""}
CONS = {"kh":"ඛ","gh":"ඝ","chh":"ඡ","jh":"ඣ","th":"ත","dh":"ද","ph":"ඵ","bh":"භ",
        "sh":"ශ","ss":"ෂ","ng":"ඟ","ny":"ඤ","t":"ට","d":"ඩ","n":"ණ","p":"ප","b":"බ",
        "m":"ම","k":"ක","g":"ග","c":"ච","j":"ජ","l":"ල","w":"ව","v":"වැ","y":"ය","r":"ර",
        "s":"ස","h":"හ","f":"ෆ","lh":"ළ","":""}
_RE_CON = re.compile("|".join(sorted(CONS, key=len, reverse=True)))
VOW_INIT = {"a":"අ","aa":"ආ","ae":"ඇ","aae":"ඈ","i":"ඉ","ii":"ඊ","u":"උ","uu":"ඌ",
            "e":"එ","ee":"ඒ","o":"ඔ","oo":"ඕ","au":"ඖ"}

def _phonetic_global(word: str) -> str:
    source_word_lower = word.lower()
    t = source_word_lower
    out_parts = [] # Use a list to build output, then join

    # Ensure logger is available or defined if used within this function like:
    # logger = logging.getLogger("SinhalaWordProcessor") # If not globally available

    while t:
        match = _RE_CON.match(t)
        c = match.group(0) if match else ""
        
        # Determine string remaining after potential consonant match
        remaining_after_c = t[len(c):]
        
        v = ""
        # If no consonant was matched (c=""), search for initial vowel in t itself.
        # If a consonant was matched (c!=""), search for a subsequent vowel in remaining_after_c.
        string_to_search_vowel_in = remaining_after_c if c else t

        if string_to_search_vowel_in: # Ensure we are not searching in an empty string for a vowel
             v = next((vv for vv in sorted(VOW, key=len, reverse=True) if string_to_search_vowel_in.startswith(vv)), "")

        if c: # Consonant matched
            out_parts.append(CONS[c])
            out_parts.append(VOW[v]) # VOW[v] is safe, as VOW[""] results in an empty string
            # Advance t: remove consumed consonant part, then consumed vowel part from that remainder
            t = remaining_after_c[len(v):]
        elif v: # No consonant, but an initial vowel matched
            out_parts.append(VOW_INIT.get(v, v)) # Use VOW_INIT for initial vowels, default to v itself
            # Advance t: remove consumed initial vowel part from original t
            t = t[len(v):] 
        else:
            # No initial consonant (c) and no initial vowel (v) were matched from the current start of t.
            # This means t[0] is an unhandled character.
            if t: # Should always be true here because of the 'while t:' condition
                # Append the unhandled character as is to the output.
                out_parts.append(t[0]) 
                # Advance t by one character to ensure progress and prevent an infinite loop.
                t = t[1:]
            else: 
                # This case should ideally not be reached if 'while t:' correctly gates the loop.
                break # Safeguard to exit if t somehow became empty here.

    result = "".join(out_parts)
    # Return the original word (as passed, not lowercased) if the phonetic conversion results in an empty string.
    # Or, stick to returning the processed (lowercased) word if output is empty.
    # The original code was `return out or word` (original case word).
    # Let's be consistent: if phonetic is empty, return original input `word`.
    return result if result else word

# ------------------------------------------------------------------
#  Main window class
# ------------------------------------------------------------------

# Import constants
from ui.constants import (
    MIN_KB_FONT, MAX_KB_FONT, BASE_KB_HEIGHT, BASE_KB_FONT, 
    DEFAULT_KB_FONT_SIZE, DEFAULT_FONT_SIZE
)

class SinhalaWordApp(QMainWindow):
    """
    SinhalaWordApp: Main application class for the Sinhala Word Processor.

    Features:
    - Realtime Singlish to Sinhala transliteration.
    - Suggestion popup (up to 9 matches, Tab/1-9/Click to accept)
    - Basic spell‑checker — unknown Sinhala words are underlined red
    - Two classic toolbars (Standard & Formatting) ‑or‑ hide them and use the menu
    - Dark / Light theme toggle (View → Toggle Theme) - Enhanced implementation
    - Plain‑text load/save functionality.
    """
    # We don't need the resize state anymore with QSplitter
    def __init__(self):
        # Call the init methods in the correct order
        self.init_preferences()
        self.init_font_manager()
        self.init_core_attributes()
        self.init_theme_manager()
        self.init_window_setup() # Calls super().__init__()
        self.init_core_ui_widgets() # Initializes self.editor
        self.init_suggestion_handling() # Now safe to create SuggestionPopup
        self.init_status_bar()
        self.init_dictionaries_and_modules()
        self.init_actions_and_shortcuts()
        self.init_keyboard()
        self.init_ui_layout_and_theme()
        self.init_event_handlers_and_timers()
        self.init_menu_and_toolbars()

    def init_preferences(self):
        """Initialize user preferences."""
        self.preferences = config.load_user_preferences()

    def init_font_manager(self):
        """Initialize the font manager and base font."""
        self.font_manager = FontManager()
        self.font_manager.update_from_preferences(self.preferences)
        self.sinhala_font_families = self.font_manager.get_available_fonts()
        self.base_font = self.font_manager.get_font()

    def init_core_attributes(self):
        """Initialize core application attributes."""
        self.MAIN_LEXICON = {}
        self.USER_MAP = {}
        self.MAP = {}
        self.USER_MAP_FP = config.USER_DICT_FILE
        self.SAVE_PENDING = False

    def init_suggestion_handling(self):
        """Initialize suggestion handling mechanisms."""
        self.buffer = []
        self.word_start_pos = None  # Position in the document where the current word started
        self.current_suggestions = []  # Store current suggestions for fixed area

        # Initialize suggestion timer
        self._suggestion_timer = QTimer()
        self._suggestion_timer.setSingleShot(True)
        self._suggestion_timer.timeout.connect(self.update_suggestion_area)

        # --- Suggestion Popup (Near Cursor) ---
        # Create a popup widget for suggestions
        self.suggestion_popup = SuggestionPopup(self)
        self.suggestion_popup.suggestionSelected.connect(self.accept_suggestion)

        # Hide the popup initially
        self.suggestion_popup.hide()

        # Log that we've created the suggestion area
        logger.info("Created suggestion label")

    def init_theme_manager(self):
        """Initialize the theme manager."""
        self.theme_manager = ThemeManager()
        # Set theme from preferences
        if self.preferences["theme"] == "dark":
            self.theme_manager.current_theme = "dark"

    def init_core_ui_widgets(self):
        """Initialize core UI widgets like the editor."""
        self.editor = QTextEdit()
        self.editor.setFont(self.base_font)

    def init_window_setup(self):
        """Perform basic window setup."""
        super().__init__()
        self.setWindowTitle("Sinhala Word Processor")
        self.resize(1100, 780)
        # Set application icon
        icon_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                                "resources", "splash", "sinhalaword_icon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

    def init_status_bar(self):
        """Initialize the status bar."""
        self.status = self.statusBar()
        self.lineCol = QLabel("Ln 1, Col 1")
        self.wordCount = QLabel("Words: 0")
        self.themeLbl = QLabel("☀️ Light" if self.theme_manager.current_theme == "light" else "🌙 Dark")
        self.status.addPermanentWidget(self.lineCol)
        self.status.addPermanentWidget(self.wordCount)
        self.status.addPermanentWidget(self.themeLbl)

    def init_dictionaries_and_modules(self):
        """Load dictionaries and initialize transliterator and spellchecker."""
        self._load_dictionaries()  # Call load method
        self.transliterator = SinhalaTransliterator(self.MAP)
        self.spellchecker = SinhalaSpellChecker(self.MAP)

    def init_actions_and_shortcuts(self):
        """Initialize actions, shortcuts, and toggle actions."""
        self.build_shortcuts()  # Initialize shortcuts and actions

        # --- Create Toggle Actions ---
        # These need to be created before building menus and toolbars
        self.singlish_toggle_action = QAction("Singlish: On" if self.preferences["singlish_enabled"] else "Singlish: Off", self, checkable=True)
        self.singlish_toggle_action.setChecked(self.preferences["singlish_enabled"])
        self.singlish_toggle_action.triggered.connect(self.toggle_singlish)

        self.suggestions_toggle_action = QAction("Suggestions: On" if self.preferences["show_suggestions"] else "Suggestions: Off", self, checkable=True)
        self.suggestions_toggle_action.setChecked(self.preferences["show_suggestions"])
        self.suggestions_toggle_action.triggered.connect(self.toggle_suggestions)

        self.keyboard_toggle_action = QAction("Keyboard: On" if self.preferences["show_keyboard"] else "Keyboard: Off", self, checkable=True)
        self.keyboard_toggle_action.setChecked(self.preferences["show_keyboard"])
        self.keyboard_toggle_action.triggered.connect(self.toggle_keyboard)
        self.keyboard_toggle_action.setIcon(self.create_icon("keyboard"))
        self.keyboard_toggle_action.setProperty("icon_name", "keyboard")

    def init_keyboard(self):
        """Initialize the on-screen keyboard."""
        is_dark_mode = self.theme_manager.is_dark_mode()
        keyboard_font_size = self.preferences.get("keyboard_font_size", DEFAULT_KB_FONT_SIZE)

        if keyboard_font_size < MIN_KB_FONT or keyboard_font_size > MAX_KB_FONT:
            keyboard_font_size = DEFAULT_KB_FONT_SIZE
            self.preferences["keyboard_font_size"] = keyboard_font_size
            logging.warning(f"Invalid keyboard font size detected, reset to {keyboard_font_size}")

        logging.info(f"Using keyboard font size: {keyboard_font_size}")

        try:
            self.keyboard_area = SinhalaKeyboard(parent=self, dark_mode=is_dark_mode)
            if hasattr(self.keyboard_area, 'keyPressed'):
                self.keyboard_area.keyPressed.connect(self.on_keyboard_button_clicked)

            if hasattr(self.keyboard_area, 'height_for_font'):
                keyboard_height = self.keyboard_area.height_for_font(keyboard_font_size)
                logging.info(f"Calculated keyboard height: {keyboard_height} based on font size: {keyboard_font_size}")
            else:
                from ui.constants import BASE_KB_HEIGHT, BASE_KB_FONT
                keyboard_height = int((keyboard_font_size / BASE_KB_FONT) * BASE_KB_HEIGHT)
                logging.info(f"Using fallback height calculation: {keyboard_height} for font size: {keyboard_font_size}")
            
            self.keyboard_area.resize(self.width(), keyboard_height)
            if hasattr(self.keyboard_area, 'update_buttons'):
                QTimer.singleShot(100, self.keyboard_area.update_buttons)
            logging.info(f"Created Sinhala keyboard with font size: {keyboard_font_size}")
        except Exception as e:
            logging.error(f"Error creating Sinhala keyboard: {e}")
            self.keyboard_area = QWidget(self)
            self.keyboard_area.setFixedHeight(100)
            logging.warning("Created empty widget as keyboard placeholder due to error")

    def init_ui_layout_and_theme(self):
        """Initialize UI layout, splitter, and apply theme."""
        self.apply_theme_to_all_widgets()
        self.update_combo_box_styles()
        self.update_icons_for_theme(self.theme_manager.current_theme)

        central_widget = QWidget()
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        logger.info("Created suggestion popup widget")

        self.main_splitter = QSplitter(Qt.Vertical)
        self.main_splitter.addWidget(self.editor)

        self.keyboard_container = QWidget()
        default_keyboard_height = 264
        logging.info(f"Keyboard preferences: height={self.preferences.get('keyboard_height', 'not set')}, "
                     f"font_size={self.preferences.get('keyboard_font_size', 'not set')}")
        
        screen = self.screen()
        if screen:
            screen_height = screen.availableGeometry().height()
            screen_width = screen.availableGeometry().width()
            logging.info(f"Screen dimensions: {screen_width}x{screen_height}")
            max_keyboard_height = int(screen_height * 0.25)
            saved_height = self.preferences.get("keyboard_height", default_keyboard_height)
            if saved_height <= 0 or saved_height > screen_height:
                logging.warning(f"Invalid saved keyboard height: {saved_height}, using default")
                saved_height = default_keyboard_height
            keyboard_height = min(saved_height, max_keyboard_height)
            logging.info(f"Using keyboard height: {keyboard_height} (saved: {saved_height}, max: {max_keyboard_height})")
        else:
            keyboard_height = self.preferences.get("keyboard_height", default_keyboard_height)
            logging.info(f"No screen info available, using keyboard height: {keyboard_height}")

        self.keyboard_container.setMinimumHeight(10)
        
        splitter_handle = QFrame(self.keyboard_container)
        splitter_handle.setFrameShape(QFrame.HLine)
        splitter_handle.setFrameShadow(QFrame.Sunken)
        splitter_handle.setFixedHeight(8)
        splitter_handle.setStyleSheet("""
            QFrame {
                background-color: #cccccc;
                border: 1px solid #aaaaaa;
                border-radius: 2px;
            }
            QFrame:hover {
                background-color: #bbbbbb;
                border: 1px solid #999999;
            }
        """)
        self.keyboard_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.keyboard_container.setMinimumHeight(80)
        self.keyboard_container.setObjectName("keyboard_container")

        keyboard_layout = QVBoxLayout(self.keyboard_container)
        keyboard_layout.setContentsMargins(0, 0, 0, 0)
        keyboard_layout.setSpacing(0)

        try:
            keyboard_layout.addWidget(self.keyboard_area)
        except Exception as e:
            logging.error(f"Error adding keyboard to layout: {e}")
            try:
                self.keyboard_area = QWidget(self)
                keyboard_layout.addWidget(self.keyboard_area)
            except Exception as fallback_error:
                logging.error(f"Failed to create fallback keyboard: {fallback_error}")

        self.main_splitter.addWidget(self.keyboard_container)

        if "splitter_state" in self.preferences:
            try:
                saved_sizes = self.preferences["splitter_state"][1]
                self.main_splitter.setSizes(saved_sizes)
                logging.info(f"Restored splitter sizes: {saved_sizes}")
            except Exception as e:
                self.main_splitter.setSizes([700, default_keyboard_height])
                logging.error(f"Error restoring splitter state: {e}")
        else:
            self.main_splitter.setSizes([700, default_keyboard_height])

        main_layout.addWidget(self.main_splitter)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setChildrenCollapsible(False)
        
        if hasattr(self.keyboard_area, 'setSizePolicy'):
            self.keyboard_area.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        if self.preferences["show_keyboard"]:
            self.keyboard_container.show()
        else:
            self.keyboard_container.hide()

        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(2)
        main_layout.setSizeConstraint(QVBoxLayout.SetNoConstraint)
        self.setCentralWidget(central_widget)
        central_widget.updateGeometry()

    def init_event_handlers_and_timers(self):
        """Initialize event handlers and timers."""
        self.editor.installEventFilter(self)
        self.editor.textChanged.connect(self.on_text_changed)
        self.editor.cursorPositionChanged.connect(self.update_format_actions)
        self.editor.setContextMenuPolicy(Qt.CustomContextMenu)
        self.editor.customContextMenuRequested.connect(self.show_context_menu)
        
        self.main_splitter.splitterMoved.connect(self.on_splitter_moved)
        if hasattr(self.keyboard_area, 'keyboardResized'):
            try:
                self.keyboard_area.keyboardResized.connect(self.on_keyboard_resized)
            except Exception as e:
                logger.error(f"Error connecting keyboard resize signal: {e}")
        else:
            logger.warning("Keyboard area doesn't have keyboardResized signal - skipping connection")

        self.spell_check_timer = QTimer(self)
        self.spell_check_timer.setSingleShot(True)
        self.spell_check_timer.timeout.connect(self.perform_spell_check)

    def init_menu_and_toolbars(self):
        """Build menus and toolbars."""
        self.build_menu()
        self.build_toolbars()
        self.update_status()  # Initial status
        self.recent_files_menu = None # This is set in build_menu
        self.update_recent_files_menu()


    def show_context_menu(self, position):
        """Show right-click context menu."""
        menu = self.editor.createStandardContextMenu()

        # Add separator and custom actions
        menu.addSeparator()

        # Add "Learn Selected Word" action
        learn_action = QAction("Learn Selected Word", self)
        learn_action.triggered.connect(self.learn_selected_word)
        menu.addAction(learn_action)

        # Get selected text
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()

            # Check if the selected text contains Sinhala characters
            is_sinhala = any("\u0D80" <= c <= "\u0DFF" for c in selected_text)

            if is_sinhala and not self.spellchecker.is_known_word(selected_text):
                # Add spell check suggestions if available
                suggestions = self.spellchecker.suggest_corrections(selected_text)
                if suggestions:
                    suggestion_menu = menu.addMenu("Spelling Suggestions")
                    for suggestion in suggestions:
                        action = QAction(suggestion, self)
                        action.triggered.connect(lambda checked, word=suggestion: self.replace_selected_text(word))
                        suggestion_menu.addAction(action)

        # Show the menu at the cursor position
        menu.exec_(self.editor.mapToGlobal(position))

    def replace_selected_text(self, replacement):
        """Replace selected text with the given replacement."""
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            cursor.beginEditBlock()
            cursor.removeSelectedText()
            cursor.insertText(replacement)
            cursor.endEditBlock()

    def learn_selected_word(self):
        """Learn selected Sinhala word and add to user dictionary."""
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()

            # Check if the selected text contains Sinhala characters
            is_sinhala = any("\u0D80" <= c <= "\u0DFF" for c in selected_text)

            if is_sinhala:
                # Ask for Singlish equivalent
                singlish, ok = QInputDialog.getText(
                    self, 
                    "Add to Dictionary", 
                    f"Enter Singlish for: {selected_text}"
                )

                if ok and singlish:
                    # Add to user dictionary
                    self.USER_MAP[singlish.lower()] = selected_text
                    # Update combined map
                    self.MAP[singlish.lower()] = selected_text
                    # Update transliterator and spellchecker
                    self.transliterator = SinhalaTransliterator(self.MAP)
                    self.spellchecker = SinhalaSpellChecker(self.MAP)
                    # Mark for saving
                    self.SAVE_PENDING = True
                    # Save immediately
                    self._save_map()
                    QMessageBox.information(self, "Success", "Word added to dictionary")
                    # Trigger spell check to remove any red underlines
                    self.perform_spell_check()
            else:
                QMessageBox.warning(self, "Warning", "Please select a Sinhala word")
        else:
            QMessageBox.warning(self, "Warning", "Please select a word first")

    # Removed on_suggestion_link_clicked method
    # as it is no longer needed with the new SuggestionPopup implementation
    # The SuggestionPopup class handles clicks internally and emits a signal
    
    def on_keyboard_button_clicked(self, char):
        """Handles clicks on the on-screen keyboard buttons."""
        try:
            # Simplified implementation matching SinhalaWordProcessor_enhanced.py
            if char == "Space":
                # Insert space and let the eventFilter handle the buffer and suggestions
                self.editor.insertPlainText(" ")
            elif char == "Backspace":
                # Let the editor handle the backspace
                self.editor.textCursor().deletePreviousChar()
            else:
                # For any other character, just insert it into the editor
                # The eventFilter will handle the buffer and suggestions
                self.editor.insertPlainText(char)
                
            # Log successful keyboard input
            logger.debug(f"Keyboard input: '{char}'")
        except Exception as e:
            logger.error(f"Error in on_keyboard_button_clicked for char '{char}': {e}")
            # Try a more direct approach on error
            try:
                # Try to insert the character directly
                if char != "Backspace" and char != "Space":
                    self.editor.insertPlainText(char)
                elif char == "Space":
                    self.editor.insertPlainText(" ")
            except Exception as inner_e:
                logger.error(f"Failed fallback insertion: {inner_e}")
            
            # Clear state on error
            self.reset_input_state()
            
    # Timer for debouncing keyboard resize events
    _keyboard_resize_timer = None
    _pending_keyboard_height = None
    # We don't need the manual resize flag anymore with QSplitter
    _in_resize = False  # Flag to prevent resize loops
    _last_keyboard_height = None  # Store last height for comparison

    def on_keyboard_resized(self, new_height):
        """Handle keyboard resize events to update the container height (debounced)."""
        try:
            # Check if we're in a manual resize - if so, skip automatic resizing
            if hasattr(self, '_manual_resize') and self._manual_resize:
                logger.debug(f"Ignoring keyboard resize signal during manual resize: {new_height}")
                return
                
            # Validate the height
            if new_height <= 0:
                logger.warning(f"Ignoring invalid keyboard height: {new_height}")
                return
            
            # Check if this is a duplicate resize (same height as before)
            if hasattr(self, '_last_keyboard_height') and self._last_keyboard_height is not None and abs(self._last_keyboard_height - new_height) < 5:
                logger.debug(f"Ignoring small height change: {self._last_keyboard_height} -> {new_height}")
                return
                
            # Store the pending height
            self._pending_keyboard_height = new_height

            # Initialize timer if it doesn't exist
            if self._keyboard_resize_timer is None:
                self._keyboard_resize_timer = QTimer(self)
                self._keyboard_resize_timer.setSingleShot(True)
                self._keyboard_resize_timer.timeout.connect(self._process_keyboard_resize)

            # Use a longer debounce delay (150ms) to reduce resize frequency
            # This is critical to prevent the rubber band effect
            self._keyboard_resize_timer.start(150)

            logger.debug(f"Keyboard resize signal received, pending height: {new_height}")

        except Exception as e:
            logger.error(f"Error in on_keyboard_resized: {e}")

    # Track resize loop detection
    _resize_count = 0
    _last_resize_time = 0
    
    @Slot()
    def request_keyboard_resize(self, new_height):
        """Request a keyboard resize to the specified height.
        This method can be called from the resize handle or other components.
        """
        if new_height <= 0:
            logger.warning(f"Ignoring invalid keyboard height request: {new_height}")
            return
            
        # Store the requested height and trigger the resize timer
        self._pending_keyboard_height = new_height
        self._keyboard_resize_timer.start()
        
    def apply_keyboard_height(self, new_height):
        """Apply a keyboard height immediately without using the timer.
        This is used for manual resizing to ensure the size is maintained.
        """
        if new_height <= 0:
            logger.warning(f"Ignoring invalid keyboard height: {new_height}")
            return
            
        # Set a flag to indicate this is a manual resize
        self._manual_resize = True
        
        try:
            # Calculate container height
            container_height = new_height + 10
            
            # No need to save keyboard height in preferences anymore
            logger.info(f"Manually applying keyboard height: {new_height}, container: {container_height}")
            
            # Resize the keyboard container directly
            if hasattr(self, 'keyboard_container') and self.keyboard_container.isVisible():
                self.keyboard_container.resize(self.keyboard_container.width(), container_height)
                
                # Resize the keyboard area directly
                if hasattr(self, 'keyboard_area'):
                    self.keyboard_area.resize(self.keyboard_area.width(), new_height)
                    
                    # Update the buttons to match the new size
                    if hasattr(self.keyboard_area, 'update_buttons'):
                        self.keyboard_area.update_buttons()
            
            # Force layout update
            self.updateGeometry()
            
            # Save preferences to ensure the size is remembered
            config.save_user_preferences(self.preferences)
            
        except Exception as e:
            logger.error(f"Error applying keyboard height: {e}")
        finally:
            # Clear the manual resize flag after a delay
            QTimer.singleShot(500, lambda: setattr(self, '_manual_resize', False))
        
    def _process_keyboard_resize(self):
        """Process the pending keyboard resize after the debounce timer times out."""
        try:
            # Skip processing if any of these conditions are true:
            # 1. User is manually resizing the keyboard
            # 2. Keyboard is in a resize operation
            # 3. Keyboard has manual font size flag set
            
            # User-initiated resize check removed (ResizeState enum no longer used)
                
            # Check if the keyboard is in a manual resize operation
            if hasattr(self.keyboard_area, 'resize_in_progress') and self.keyboard_area.resize_in_progress:
                logger.debug("Skipping automatic resize as keyboard is in a manual resize operation")
                return
                
            # Check if the keyboard has the manual font size flag set
            if hasattr(self.keyboard_area, '_manual_font_size') and self.keyboard_area._manual_font_size:
                logger.debug("Skipping automatic resize as keyboard has manual font size flag set")
                return
                
            # Get the pending height
            new_height = self._pending_keyboard_height
            self._pending_keyboard_height = None # Clear pending height

            if new_height is None:
                return # No pending resize
                
            # Check for resize loops with a simpler approach
            import time
            current_time = time.time()
            if hasattr(self, '_last_resize_time') and current_time - self._last_resize_time < 0.2:
                # If resizes are happening too quickly, skip this one
                logger.debug("Skipping resize - too soon after previous resize")
                return
                
            self._last_resize_time = current_time

            logger.info(f"Processing keyboard resize to height: {new_height}")

            # Validate the new height to ensure it's reasonable
            if new_height <= 0:
                logger.warning(f"Ignoring invalid processed keyboard height: {new_height}")
                return

            # Check if this is a duplicate resize (same height as before)
            if hasattr(self, '_last_keyboard_height') and abs(self._last_keyboard_height - new_height) < 5:
                logger.debug(f"Skipping resize - height change too small: {self._last_keyboard_height} -> {new_height}")
                return

            # Store the height for future comparison
            self._last_keyboard_height = new_height

            # Set a flag to indicate we're in a programmatic resize
            # This prevents resize loops when the keyboard itself is being resized
            if hasattr(self, '_in_resize') and self._in_resize:
                logger.debug("Skipping resize as we're already in a resize operation")
                return
                
            self._in_resize = True

            try:
                # Don't change minimum height here - keep it at the global minimum
                # Just resize the container to match the keyboard
                container_height = new_height + 10
    
                # Save the keyboard height in preferences
                self.preferences["keyboard_height"] = new_height
                logger.info(f"Keyboard resized to height: {new_height}, container: {container_height}")
                
                # Resize the keyboard container
                if hasattr(self, 'keyboard_container') and self.keyboard_container.isVisible():
                    self.keyboard_container.resize(self.keyboard_container.width(), container_height)
                    
                    # Resize the keyboard area directly
                    if hasattr(self, 'keyboard_area'):
                        self.keyboard_area.resize(self.keyboard_area.width(), new_height)
    
                # Check if the keyboard is in a manual resize operation
                # If so, we should respect its current font size
                keyboard_in_manual_resize = (hasattr(self.keyboard_area, '_manual_font_size') and 
                                           self.keyboard_area._manual_font_size)
                
                # Only adjust font size if not in a manual resize operation
                if not keyboard_in_manual_resize:
                    # Calculate a new font size based on the keyboard height
                    current_font_size = self.preferences.get("keyboard_font_size", BASE_KB_FONT)
                    
                    # Calculate new font size proportional to height and round to nearest integer
                    new_font_size = max(MIN_KB_FONT, min(MAX_KB_FONT, round(BASE_KB_FONT * new_height / BASE_KB_HEIGHT)))
                    
                    # Only update if the integer value has changed significantly
                    if abs(new_font_size - current_font_size) > 2:
                        logger.info(f"Adjusting keyboard font size from {current_font_size} to {new_font_size} based on height {new_height}")
                        self.preferences["keyboard_font_size"] = new_font_size
                        
                        # Update the keyboard font size using the proper method
                        if hasattr(self.keyboard_area, 'set_font_size'):
                            self.keyboard_area.set_font_size(new_font_size)
                else:
                    # We're in a manual resize operation, so respect the keyboard's current font size
                    logger.info(f"Skipping font size adjustment during manual keyboard resize")
                    
                    # Update the preferences with the keyboard's current font size
                    if hasattr(self.keyboard_area, 'font_size'):
                        current_font_size = self.keyboard_area.font_size
                        self.preferences["keyboard_font_size"] = current_font_size
                        logger.info(f"Updated preferences with current keyboard font size: {current_font_size}")
                    
                # Force layout update to apply changes
                self.updateGeometry()
            except Exception as layout_error:
                logger.error(f"Error updating layout after processed keyboard resize: {layout_error}")
            finally:
                # Always clear the resize flag
                self._in_resize = False

        except Exception as e:
            logger.error(f"Error in _process_keyboard_resize: {e}")
            # Try to recover by forcing a layout update
            try:
                self.updateGeometry()
            except:
                pass


    def ensure_keyboard_fits_screen(self):
        """Ensure the keyboard size fits within the available screen space"""
        try:
            # Check if we're in a manual resize - if so, skip automatic adjustments
            if hasattr(self, '_manual_resize') and self._manual_resize:
                logger.debug("Skipping keyboard size check as we're in a manual resize operation")
                return
                
            # Also check if the keyboard is in a manual resize operation
            if hasattr(self.keyboard_area, 'resize_in_progress') and self.keyboard_area.resize_in_progress:
                logger.debug("Skipping keyboard size check as keyboard is in a manual resize operation")
                return
                
            # Also check if the keyboard has the manual font size flag set
            if hasattr(self.keyboard_area, '_manual_font_size') and self.keyboard_area._manual_font_size:
                logger.debug("Skipping keyboard size check as keyboard has manual font size flag set")
                return
                
            # Get current screen size
            screen = self.screen()
            if not screen:
                logger.warning("Could not get screen information")
                return

            screen_height = screen.availableGeometry().height()
            screen_width = screen.availableGeometry().width()
            
            # Log screen dimensions for debugging
            logger.info(f"Screen dimensions: {screen_width}x{screen_height}")

            # Get current window size
            window_height = self.height()
            window_width = self.width()

            # Get current keyboard height with error handling
            try:
                # Get the current keyboard height
                keyboard_height = self.keyboard_area.height()
                
                # If the keyboard height is invalid, calculate based on font size
                if keyboard_height <= 0:
                    logger.debug(f"Invalid keyboard height: {keyboard_height}, calculating from font size")
                    keyboard_font_size = self.preferences.get("keyboard_font_size", DEFAULT_KB_FONT_SIZE)
                    if hasattr(self.keyboard_area, 'height_for_font'):
                        keyboard_height = self.keyboard_area.height_for_font(keyboard_font_size)
                    else:
                        # Fallback calculation if the method is not available
                        from ui.constants import BASE_KB_HEIGHT, BASE_KB_FONT
                        keyboard_height = int((keyboard_font_size / BASE_KB_FONT) * BASE_KB_HEIGHT)
            except Exception as e:
                logger.error(f"Error getting keyboard height: {e}")
                keyboard_font_size = self.preferences.get("keyboard_font_size", DEFAULT_KB_FONT_SIZE)
                if hasattr(self.keyboard_area, 'height_for_font'):
                    keyboard_height = self.keyboard_area.height_for_font(keyboard_font_size)
                else:
                    # Fallback calculation if the method is not available
                    from ui.constants import BASE_KB_HEIGHT, BASE_KB_FONT
                    keyboard_height = int((keyboard_font_size / BASE_KB_FONT) * BASE_KB_HEIGHT)

            # Calculate maximum allowed keyboard height (e.g., 50% of screen height)
            max_keyboard_height = int(screen_height * 0.50)
            
            # Log the keyboard height information
            logger.info(f"Using keyboard height: {keyboard_height} (saved: {self.preferences.get('keyboard_height', 'not set')}, max: {max_keyboard_height})")

            # If keyboard is too large, resize it
            if keyboard_height > max_keyboard_height:
                # Only log at debug level to reduce noise during normal operation
                logger.debug(f"Limiting keyboard height to {max_keyboard_height} (50% of screen height)")

                # Check if we're already at the max height (within a small margin)
                # to avoid triggering unnecessary resize events
                if abs(keyboard_height - max_keyboard_height) > 5:
                    try:
                        # Resize keyboard without setting fixed height
                        self.keyboard_area.resize(self.keyboard_area.width(), max_keyboard_height)
                        self.keyboard_container.resize(self.keyboard_container.width(), max_keyboard_height + 10)
                        
                        # No need to store keyboard height in preferences anymore

                        # Update buttons to match new size
                        if hasattr(self.keyboard_area, 'update_buttons'):
                             self.keyboard_area.update_buttons()

                        # Force layout update to ensure changes take effect
                        self.keyboard_area.updateGeometry()
                        self.keyboard_container.updateGeometry()
                    except Exception as resize_error:
                        logger.error(f"Error resizing keyboard: {resize_error}")
            else:
                # If the keyboard height is valid and not too large, respect it
                # This is important for manual resizing
                current_container_height = self.keyboard_container.height()
                expected_container_height = keyboard_height + 10
                
                # Only adjust if there's a significant difference
                if abs(current_container_height - expected_container_height) > 5:
                    logger.debug(f"Adjusting container height to match keyboard: {expected_container_height}")
                    self.keyboard_container.resize(self.keyboard_container.width(), expected_container_height)

            # Now ensure the entire window fits on screen
            try:
                self.ensure_window_fits_screen()
            except Exception as window_error:
                logger.error(f"Error ensuring window fits screen: {window_error}")

            # Log screen and window info at debug level to reduce noise
            logger.debug(f"Screen: {screen_width}x{screen_height}, Window: {self.width()}x{self.height()}, Keyboard height: {self.keyboard_area.height()}")
        except Exception as e:
            logger.error(f"Error in ensure_keyboard_fits_screen: {e}")
            # Try to recover by forcing a layout update
            try:
                self.updateGeometry()
            except:
                pass

    def ensure_window_fits_screen(self):
        """Ensure the entire application window fits within the screen"""
        try:
            # Get current screen size
            screen = self.screen()
            if not screen:
                logger.warning("Could not get screen information")
                return

            screen_rect = screen.availableGeometry()
            screen_height = screen_rect.height()
            screen_width = screen_rect.width()

            # Get current window size
            window_height = self.height()
            window_width = self.width()

            # Calculate maximum allowed window size (90% of screen)
            max_window_height = int(screen_height * 0.9)
            max_window_width = int(screen_width * 0.9)

            # If window is too large, resize it
            resized = False
            if window_height > max_window_height:
                logger.info(f"Adjusting window height from {window_height} to {max_window_height} to fit screen")
                window_height = max_window_height
                resized = True

            if window_width > max_window_width:
                logger.info(f"Adjusting window width from {window_width} to {max_window_width} to fit screen")
                window_width = max_window_width
                resized = True

            if resized:
                self.resize(window_width, window_height)

            # Also ensure the window is positioned on screen
            window_pos = self.pos()
            if not screen_rect.contains(self.frameGeometry()):
                # Center the window on screen
                new_x = screen_rect.center().x() - window_width // 2
                new_y = screen_rect.center().y() - window_height // 2
                self.move(new_x, new_y)
                logger.info(f"Repositioned window to center of screen")

            logger.info(f"Final window size: {self.width()}x{self.height()}")
        except Exception as e:
            logger.error(f"Error in ensure_window_fits_screen: {e}")

    def create_icon(self, name):
        """Create an icon for toolbar buttons using pyside_icons.py"""
        try:
            # Import the get_toolbar_icon function from ui.icons.py
            from ui.icons import get_toolbar_icon
            # Use the current theme to get the appropriate icon color
            theme = "dark" if self.theme_manager.is_dark_mode() else "light"
            return get_toolbar_icon(name, theme=theme)
        except (ImportError, AttributeError):
            # Fallback if ui.icons.py is not available or function not found
            return None

    def build_shortcuts(self):
        """Sets up keyboard shortcuts for common actions."""
        # File actions
        self.new_action = QAction("New", self)
        self.new_action.setShortcut("Ctrl+N")
        self.new_action.triggered.connect(self.new_file)
        self.new_action.setIcon(self.create_icon("new"))
        self.new_action.setProperty("icon_name", "new")

        self.open_action = QAction("Open...", self)
        self.open_action.setShortcut("Ctrl+O")
        self.open_action.triggered.connect(self.open_file)
        self.open_action.setIcon(self.create_icon("open"))
        self.open_action.setProperty("icon_name", "open")

        self.save_action = QAction("Save", self)
        self.save_action.setShortcut("Ctrl+S")
        self.save_action.triggered.connect(self.save_file)
        self.save_action.setIcon(self.create_icon("save"))
        self.save_action.setProperty("icon_name", "save")

        # Edit actions
        self.cut_action = QAction("Cut", self)
        self.cut_action.setShortcut("Ctrl+X")
        self.cut_action.triggered.connect(self.editor.cut)
        self.cut_action.setIcon(self.create_icon("cut"))
        self.cut_action.setProperty("icon_name", "cut")

        self.copy_action = QAction("Copy", self)
        self.copy_action.setShortcut("Ctrl+C")
        self.copy_action.triggered.connect(self.editor.copy)
        self.copy_action.setIcon(self.create_icon("copy"))
        self.copy_action.setProperty("icon_name", "copy")

        self.paste_action = QAction("Paste", self)
        self.paste_action.setShortcut("Ctrl+V")
        self.paste_action.triggered.connect(self.editor.paste)
        self.paste_action.setIcon(self.create_icon("paste"))
        self.paste_action.setProperty("icon_name", "paste")

        self.undo_action = QAction("Undo", self)
        self.undo_action.setShortcut("Ctrl+Z")
        self.undo_action.triggered.connect(self.editor.undo)
        self.undo_action.setIcon(self.create_icon("undo"))
        self.undo_action.setProperty("icon_name", "undo")

        self.redo_action = QAction("Redo", self)
        self.redo_action.setShortcut("Ctrl+Y")
        self.redo_action.triggered.connect(self.editor.redo)
        self.redo_action.setIcon(self.create_icon("redo"))
        self.redo_action.setProperty("icon_name", "redo")

        # View actions
        self.toggle_theme_action = QAction("Toggle Theme", self)
        self.toggle_theme_action.setShortcut("Ctrl+T")
        self.toggle_theme_action.triggered.connect(self.toggle_theme)
        self.toggle_theme_action.setIcon(self.create_icon("theme"))
        self.toggle_theme_action.setProperty("icon_name", "theme")

        # Add actions to the main window (this makes them available globally)
        self.addAction(self.new_action)
        self.addAction(self.open_action)
        self.addAction(self.save_action)
        self.addAction(self.cut_action)
        self.addAction(self.copy_action)
        self.addAction(self.paste_action)
        self.addAction(self.undo_action)
        self.addAction(self.redo_action)
        self.addAction(self.toggle_theme_action)

        # Print actions
        self.print_action = QAction("Print...", self)
        self.print_action.setShortcut("Ctrl+P")
        self.print_action.triggered.connect(self.print_document)
        self.print_action.setIcon(self.create_icon("print"))
        self.print_action.setProperty("icon_name", "print")
        self.addAction(self.print_action)

        self.print_preview_action = QAction("Print Preview...", self)
        # No default shortcut for print preview, but can be added if desired
        self.print_preview_action.triggered.connect(self.print_preview_document)
        self.print_preview_action.setIcon(self.create_icon("print-preview"))
        self.print_preview_action.setProperty("icon_name", "print-preview")
        self.addAction(self.print_preview_action)

# This is a duplicate method - removing it

    def toggle_toolbars(self):
        """Toggle visibility of toolbars."""
        # Get all toolbars
        toolbars = [toolbar for toolbar in self.findChildren(QToolBar)]

        # Check if any toolbar is visible
        any_visible = any(toolbar.isVisible() for toolbar in toolbars)

        # Toggle visibility
        for toolbar in toolbars:
            toolbar.setVisible(not any_visible)

    def toggle_keyboard(self):
        """Toggle on-screen keyboard visibility."""
        # Update action text based on checked state
        if self.keyboard_toggle_action.isChecked():
            self.keyboard_toggle_action.setText("Keyboard: On")
            self.keyboard_container.show()
            self.keyboard_area.show()
            # Update preferences
            self.preferences["show_keyboard"] = True
        else:
            self.keyboard_toggle_action.setText("Keyboard: Off")
            self.keyboard_container.hide()
            self.keyboard_area.hide()
            # Update preferences
            self.preferences["show_keyboard"] = False
    
    def detach_keyboard(self):
        """Detach the keyboard to make it a floating window."""
        try:
            # Check if the keyboard is a SinhalaKeyboard
            if hasattr(self.keyboard_area, 'make_detachable'):
                # Check if keyboard is already detached
                if hasattr(self, '_keyboard_dialog') and self._keyboard_dialog is not None:
                    logger.info("Keyboard is already detached")
                    # Bring the existing dialog to front
                    self._keyboard_dialog.raise_()
                    self._keyboard_dialog.activateWindow()
                    return
                
                # Hide the keyboard container
                self.keyboard_container.hide()
                
                # Make the keyboard detachable
                dialog = self.keyboard_area.make_detachable()
                
                # Store the dialog reference to prevent garbage collection
                self._keyboard_dialog = dialog
                
                # Update preferences
                self.preferences["keyboard_detached"] = True
                
                # Save preferences immediately to disk
                from app import config
                config.save_user_preferences(self.preferences)
                
                logger.info("Keyboard detached to floating window")
            else:
                logger.warning("Cannot detach keyboard - not a SinhalaKeyboard instance")
        except Exception as e:
            logger.error(f"Error detaching keyboard: {e}")
            
    def dock_keyboard(self):
        """Dock the keyboard back into the main window."""
        try:
            # Check if we have a detached keyboard dialog
            if hasattr(self, '_keyboard_dialog') and self._keyboard_dialog is not None:
                # Get the keyboard from the dialog
                if hasattr(self.keyboard_area, 'make_embedded'):
                    # Re-embed the keyboard
                    self.keyboard_area.make_embedded(self, self._keyboard_dialog)
                    
                    # Clear the dialog reference
                    self._keyboard_dialog = None
                    
                    logger.info("Keyboard docked back to main window")
                else:
                    logger.warning("Cannot dock keyboard - not a SinhalaKeyboard instance")
            else:
                logger.info("No detached keyboard to dock")
        except Exception as e:
            logger.error(f"Error docking keyboard: {e}")
            
    def reset_keyboard_size(self):
        """Reset the keyboard to a specific size (font size 20, height 225px)."""
        try:
            # Set specific font size and height as requested
            custom_font_size = 20
            custom_height = 225
            
            # First set the font size
            if hasattr(self.keyboard_area, 'set_font_size'):
                self.keyboard_area.set_font_size(custom_font_size)
                self.preferences["keyboard_font_size"] = custom_font_size
                logger.info(f"Set keyboard font size to: {custom_font_size}")
            
            # Update preferences
            self.preferences["keyboard_height"] = custom_height
            logger.info(f"Set keyboard height preference to: {custom_height}")
            
            # If we have a splitter, adjust the sizes
            if hasattr(self, 'main_splitter') and self.main_splitter:
                # Get current sizes
                current_sizes = self.main_splitter.sizes()
                if len(current_sizes) >= 2:
                    # Calculate new sizes
                    total_height = sum(current_sizes)
                    editor_height = total_height - custom_height
                    
                    # Set new sizes
                    self.main_splitter.setSizes([editor_height, custom_height])
                    logger.info(f"Adjusted splitter sizes: editor={editor_height}, keyboard={custom_height}")
            
            # Also directly resize the keyboard and container for good measure
            self.keyboard_area.resize(self.keyboard_area.width(), custom_height)
            self.keyboard_container.resize(self.keyboard_container.width(), custom_height + 10)
            
            # Update buttons to match new size
            if hasattr(self.keyboard_area, 'update_buttons'):
                self.keyboard_area.update_buttons()
                logger.info("Updated keyboard buttons")
            
            # Force layout update
            self.keyboard_container.updateGeometry()
            self.keyboard_area.updateGeometry()
            self.updateGeometry()
            
            logger.info(f"Keyboard reset to height: {custom_height} with font size: {custom_font_size}")
        except Exception as e:
            logger.error(f"Error in reset_keyboard_size: {e}")

        # Save preferences
        config.save_user_preferences(self.preferences)
        
    def set_keyboard_font_size(self, size):
        """Update the keyboard font size."""
        try:
            # Convert to integer for consistency
            size = round(float(size))
            
            # Validate size
            if size < MIN_KB_FONT or size > MAX_KB_FONT:
                logger.warning(f"Invalid keyboard font size: {size}")
                return
                
            # Update font manager - this is the single source of truth
            # The FontManager will emit a signal that the keyboard will listen to
            self.font_manager.set_keyboard_font_size(size)
            
            # Update preferences
            self.preferences["keyboard_font_size"] = size
            
            # No need to update keyboard directly - it will respond to the signal
            logger.info(f"Keyboard font size updated to: {size} via FontManager")
            
            # Save preferences
            config.save_user_preferences(self.preferences)
        except Exception as e:
            logger.error(f"Error setting keyboard font size: {e}")
    
    def _reset_keyboard_font_flag(self):
        """Reset the manual font size flag after a delay."""
        try:
            if hasattr(self.keyboard_area, '_manual_font_size'):
                self.keyboard_area._manual_font_size = False
                logger.info("Reset keyboard manual font size flag")
        except Exception as e:
            logger.error(f"Error resetting keyboard font flag: {e}")
        
    def change_font_family(self, font_name):
        """Change the font family for the editor."""
        if not font_name:
            return
            
        # Create a new font with the selected family and current size
        current_size = self.base_font.pointSize()
        new_font = QFont(font_name, current_size)
        
        # Apply to editor
        self.editor.setFont(new_font)
        
        # Apply to suggestion area
        self.suggestion_area.setFont(new_font)
        
        # Update base font
        self.base_font = new_font
        
        # Update preferences
        self.preferences["font"] = font_name
        config.save_user_preferences(self.preferences)
        
    def change_editor_font_size(self, size_text):
        """Change the font size for the editor."""
        try:
            # Convert size text to float
            size = float(size_text)
            if size <= 0:
                return
                
            # Create a new font with the current family and selected size
            font_name = self.base_font.family()
            new_font = QFont(font_name, size)
            
            # Apply to editor
            self.editor.setFont(new_font)
            
            # Apply to suggestion area
            self.suggestion_area.setFont(new_font)
            
            # Update base font
            self.base_font = new_font
            
            # Update preferences
            self.preferences["font_size"] = size
            config.save_user_preferences(self.preferences)
        except (ValueError, TypeError):
            # Ignore invalid input
            pass

    def toggle_singlish(self):
        """Toggle Singlish transliteration."""
        # Update action text based on checked state
        if self.singlish_toggle_action.isChecked():
            self.singlish_toggle_action.setText("Singlish: On")
            # Enable Singlish functionality
            # No need to do anything else as the event filter will handle it
            self.preferences["singlish_enabled"] = True
        else:
            self.singlish_toggle_action.setText("Singlish: Off")
            # Clear any pending buffer
            self.buffer.clear()
            self.word_start_pos = None
            self.clear_suggestion_area()
            self.preferences["singlish_enabled"] = False

        # Save preferences
        config.save_user_preferences(self.preferences)

    def toggle_suggestions(self):
        """Toggle suggestions display."""
        # Update action text based on checked state
        if self.suggestions_toggle_action.isChecked():
            self.suggestions_toggle_action.setText("Suggestions: On")
            # Show suggestions if there are any in the buffer
            if self.buffer:
                self.update_suggestion_area()
            self.preferences["show_suggestions"] = True
            logger.info("Suggestions enabled")
        else:
            self.suggestions_toggle_action.setText("Suggestions: Off")
            # Hide suggestion popup
            self.clear_suggestion_area()
            self.preferences["show_suggestions"] = False
            logger.info("Suggestions disabled")

        # Save preferences
        config.save_user_preferences(self.preferences)

    def build_toolbars(self):
        """Creates toolbars for standard and formatting actions."""
        # Standard Toolbar
        self.standard_toolbar = self.addToolBar("Standard")
        self.standard_toolbar.setObjectName("StandardToolbar")

        # Add actions to the standard toolbar
        self.standard_toolbar.addAction(self.new_action)
        self.standard_toolbar.addAction(self.open_action)
        self.standard_toolbar.addAction(self.save_action)
        self.standard_toolbar.addAction(self.print_action)
        self.standard_toolbar.addAction(self.print_preview_action)
        self.standard_toolbar.addSeparator()
        self.standard_toolbar.addAction(self.cut_action)
        self.standard_toolbar.addAction(self.copy_action)
        self.standard_toolbar.addAction(self.paste_action)
        self.standard_toolbar.addSeparator()
        self.standard_toolbar.addAction(self.undo_action)
        self.standard_toolbar.addAction(self.redo_action)

        # Formatting Toolbar
        self.formatting_toolbar = self.addToolBar("Formatting")
        self.formatting_toolbar.setObjectName("FormattingToolbar")

        # Custom Font ComboBox for Sinhala fonts
        self.font_combo = QComboBox(self)
        
        # Add Sinhala fonts from our fonts directory
        if self.sinhala_font_families:
            self.font_combo.addItems(self.sinhala_font_families)
        else:
            # Fallback to system fonts that support Sinhala
            default_fonts = ["Iskoola Pota", "Nirmala UI", "Latha"]
            for font in default_fonts:
                self.font_combo.addItem(font)
        
        # Set current font from preferences or use first available Sinhala font
        current_font_name = self.preferences["font"]
        current_font_index = self.font_combo.findText(current_font_name)
        if current_font_index >= 0:
            self.font_combo.setCurrentIndex(current_font_index)
        
        # Connect font change signal to update editor font
        self.font_combo.currentTextChanged.connect(self.change_font_family)
        self.formatting_toolbar.addWidget(self.font_combo)

        # Font Size ComboBox
        self.size_combo = QComboBox(self)
        self.size_combo.setEditable(True)
        self.size_combo.setFixedWidth(50)  # Slightly wider to ensure numbers are visible
        self.size_combo.setMaxVisibleItems(10)  # Limit visible items to prevent excessive scrolling
        self.size_combo.setMaxCount(len(config.FONT_SIZES))  # Limit total items
        
        # Set alignment for the line edit inside the combo box
        if self.size_combo.lineEdit():
            self.size_combo.lineEdit().setAlignment(Qt.AlignCenter)
        
        # Populate with common font sizes
        sizes = config.FONT_SIZES
        self.size_combo.addItems([str(s) for s in sizes])
        
        # Set default size
        current_size = self.base_font.pointSize()
        default_size_index = sizes.index(current_size) if current_size in sizes else -1
        if default_size_index != -1:
             self.size_combo.setCurrentIndex(default_size_index)
        else:
             self.size_combo.setCurrentText(str(current_size))

        # Add custom styling for better appearance
        self.size_combo.setStyleSheet("""
            QComboBox { 
                padding-right: 15px;
                padding-left: 5px;
                font-weight: bold;  /* Make the font bold for better visibility */
            }
            QComboBox::item {
                padding: 3px;
                font-weight: bold;  /* Make dropdown items bold too */
            }
        """)
        
        # Apply custom delegate for better rendering
        self.size_combo.setItemDelegate(FontSizeDelegate(self.size_combo))

        # Connect size change signals
        self.size_combo.currentTextChanged.connect(self.change_font_size)  # Keep using the more comprehensive version
        self.formatting_toolbar.addWidget(self.size_combo)
        
        # Add separator
        self.formatting_toolbar.addSeparator()
        
        # Add text formatting actions (bold, italic, underline)
        self.bold_action = QAction("Bold", self)
        self.bold_action.setCheckable(True)
        self.bold_action.setIcon(self.create_icon("bold"))
        self.bold_action.setProperty("icon_name", "bold")
        self.bold_action.setShortcut("Ctrl+B")
        self.bold_action.triggered.connect(self.toggle_bold)
        self.formatting_toolbar.addAction(self.bold_action)
        
        self.italic_action = QAction("Italic", self)
        self.italic_action.setCheckable(True)
        self.italic_action.setIcon(self.create_icon("italic"))
        self.italic_action.setProperty("icon_name", "italic")
        self.italic_action.setShortcut("Ctrl+I")
        self.italic_action.triggered.connect(self.toggle_italic)
        self.formatting_toolbar.addAction(self.italic_action)
        
        self.underline_action = QAction("Underline", self)
        self.underline_action.setCheckable(True)
        self.underline_action.setIcon(self.create_icon("underline"))
        self.underline_action.setProperty("icon_name", "underline")
        self.underline_action.setShortcut("Ctrl+U")
        self.underline_action.triggered.connect(self.toggle_underline)
        self.formatting_toolbar.addAction(self.underline_action)

        # --- Feature Toggles Toolbar ---
        self.toggles_toolbar = self.addToolBar("Features")
        self.toggles_toolbar.setObjectName("FeaturesToolbar")

        # Set tool button style to show both icon and text
        self.toggles_toolbar.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        
        # Add actions to the toggles toolbar
        self.toggles_toolbar.addAction(self.singlish_toggle_action)
        self.toggles_toolbar.addAction(self.keyboard_toggle_action)  # Use the action directly like other toggles
        self.toggles_toolbar.addAction(self.suggestions_toggle_action)
        self.toggles_toolbar.addSeparator()
        self.toggles_toolbar.addAction(self.toggle_theme_action)
        
    def change_font_size(self, size_text):
        """Change the font size of the selected text or the entire editor if no selection."""
        try:
            # Convert the size text to an integer
            size = int(size_text)
            
            # Get the current cursor
            cursor = self.editor.textCursor()
            
            if cursor.hasSelection():
                # Apply font size only to selected text
                format = QTextCharFormat()
                format.setFontPointSize(size)
                cursor.mergeCharFormat(format)
                self.editor.setTextCursor(cursor)
            else:
                # No selection, update the default font for future text
                self.base_font.setPointSize(size)
                self.editor.setFont(self.base_font)
                # Font for suggestion popup is handled internally
                
                # Update preferences
                self.preferences["font_size"] = size
                config.save_user_preferences(self.preferences)
            
            # Log the change
            logging.info(f"Font size changed to {size}")
        except ValueError:
            # Handle invalid input
            logging.warning(f"Invalid font size: {size_text}")
            
            # Reset to current size
            current_size = self.base_font.pointSize()
            self.size_combo.setCurrentText(str(current_size))
            
    def change_font_family(self, font_name):
        """Change the font family of the selected text or the entire editor if no selection."""
        # Get the current cursor
        cursor = self.editor.textCursor()
        
        if cursor.hasSelection():
            # Apply font family only to selected text
            format = QTextCharFormat()
            format.setFontFamily(font_name)
            cursor.mergeCharFormat(format)
            self.editor.setTextCursor(cursor)
        else:
            # No selection, update the default font for future text
            self.base_font.setFamily(font_name)
            self.editor.setFont(self.base_font)
            # Font for suggestion popup is handled internally
            
            # Update preferences
            self.preferences["font"] = font_name
            config.save_user_preferences(self.preferences)
        
        # Log the change
        logging.info(f"Font family changed to {font_name}")
        
    def toggle_bold(self):
        """Toggle bold formatting for selected text."""
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            format = QTextCharFormat()
            # Toggle bold state
            if self.bold_action.isChecked():
                format.setFontWeight(QFont.Bold)
            else:
                format.setFontWeight(QFont.Normal)
            cursor.mergeCharFormat(format)
            self.editor.setTextCursor(cursor)
        else:
            # If no selection, set the state for future typing
            format = self.editor.currentCharFormat()
            if self.bold_action.isChecked():
                format.setFontWeight(QFont.Bold)
            else:
                format.setFontWeight(QFont.Normal)
            self.editor.setCurrentCharFormat(format)
    
    def toggle_italic(self):
        """Toggle italic formatting for selected text."""
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            format = QTextCharFormat()
            format.setFontItalic(self.italic_action.isChecked())
            cursor.mergeCharFormat(format)
            self.editor.setTextCursor(cursor)
        else:
            # If no selection, set the state for future typing
            format = self.editor.currentCharFormat()
            format.setFontItalic(self.italic_action.isChecked())
            self.editor.setCurrentCharFormat(format)
    
    def toggle_underline(self):
        """Toggle underline formatting for selected text."""
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            fmt = QTextCharFormat()
            # Use the same explicit flag toggle approach
            fmt.setFontUnderline(not cursor.charFormat().fontUnderline())
            cursor.mergeCharFormat(fmt)
            self.editor.setTextCursor(cursor)
        else:
            # If no selection, set the state for future typing
            fmt = self.editor.currentCharFormat()
            # Use the same explicit flag toggle approach
            fmt.setFontUnderline(not fmt.fontUnderline())
            self.editor.setCurrentCharFormat(fmt)
    
    def update_format_actions(self):
        """Update the state of formatting actions based on current cursor position."""
        cursor = self.editor.textCursor()
        format = cursor.charFormat()
        
        # Update bold action
        self.bold_action.setChecked(format.fontWeight() == QFont.Bold)
        
        # Update italic action
        self.italic_action.setChecked(format.fontItalic())
        
        # Update underline action
        self.underline_action.setChecked(format.fontUnderline())
    
    def update_combo_box_styles(self):
        """Update combo box styles based on current theme."""
        is_dark = self.theme_manager.is_dark_mode()
        
        # Get colors from theme manager
        if is_dark:
            dropdown_bg = self.theme_manager.get_color("SecondaryBackgroundColor")
            dropdown_fg = self.theme_manager.get_color("PrimaryForegroundColor")
            dropdown_border = self.theme_manager.get_color("StrokeColor")
            selection_bg = self.theme_manager.get_color("SelectedColor")
            selection_fg = self.theme_manager.get_color("PrimaryForegroundColor")
        else:
            dropdown_bg = self.theme_manager.get_color("PrimarySolidBackgroundColor")
            dropdown_fg = self.theme_manager.get_color("PrimaryForegroundColor")
            dropdown_border = self.theme_manager.get_color("StrokeColor")
            selection_bg = self.theme_manager.get_color("SelectedColor")
            selection_fg = self.theme_manager.get_color("PrimaryForegroundColor")
        
        # Common combo box style
        combo_style = f"""
            QComboBox {{ 
                background-color: {dropdown_bg};
                color: {dropdown_fg};
                border: 1px solid {dropdown_border};
                padding-right: 12px;
                padding-left: 2px;
            }}
            QComboBox::drop-down {{
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;
                border-left: 1px solid {dropdown_border};
            }}
            QComboBox::item {{
                padding: 3px;
                background-color: {dropdown_bg};
                color: {dropdown_fg};
            }}
            QComboBox QAbstractItemView {{
                background-color: {dropdown_bg};
                color: {dropdown_fg};
                border: 1px solid {dropdown_border};
                selection-background-color: {selection_bg};
                selection-color: {selection_fg};
                outline: 0px;
            }}
        """
        
        # Apply styles to font size combo box
        if hasattr(self, 'size_combo'):
            self.size_combo.setStyleSheet(combo_style + """
                QComboBox { 
                    text-align: center;
                }
            """)
            
        # Apply styles to font combo box
        if hasattr(self, 'font_combo'):
            self.font_combo.setStyleSheet(combo_style)
            
    def apply_theme_to_all_widgets(self):
        """Apply the current theme to all widgets for consistent styling."""
        # Get the current theme
        theme = self.theme_manager.current_theme
        is_dark = theme == "dark"
        stylesheet = self.theme_manager.get_stylesheet()
        
        # Apply main stylesheet
        self.setStyleSheet(stylesheet)
        
        # Update theme label
        self.themeLbl.setText("☀️ Light" if not is_dark else "🌙 Dark")
        
        # Update keyboard theme if it's a SinhalaKeyboard
        if hasattr(self.keyboard_area, 'set_dark_mode'):
            self.keyboard_area.set_dark_mode(is_dark)
        else:
            logging.warning("Keyboard area doesn't have set_dark_mode method - using fallback styling")
            # Apply fallback styling for the basic QWidget
            if is_dark:
                self.keyboard_area.setStyleSheet("background-color: #2d2d2d; border: 1px solid #444444;")
            else:
                self.keyboard_area.setStyleSheet("background-color: #f0f0f0; border: 1px solid #cccccc;")
        
        # Update combo box styles
        self.update_combo_box_styles()
        
        # Update icons for the current theme
        self.update_icons_for_theme(theme)
        
        # Update suggestion popup theme
        theme_colors = {
            "OverlayColor": self.theme_manager.get_color("OverlayColor"),
            "PrimarySolidBackgroundColor": self.theme_manager.get_color("PrimarySolidBackgroundColor"),
            "PrimarySolidBorderColor": self.theme_manager.get_color("PrimarySolidBorderColor"),
            "PrimaryForegroundColor": self.theme_manager.get_color("PrimaryForegroundColor"),
            "MouseOverBackgroundColor": self.theme_manager.get_color("MouseOverBackgroundColor"),
            "AccentColor": self.theme_manager.get_color("AccentColor"),
            "AccentPressedColor": self.theme_manager.get_color("AccentPressedColor")
        }
        self.suggestion_popup.update_theme(is_dark, theme_colors)
        
        # Apply theme to all child widgets recursively
        self.apply_theme_to_widget(self, is_dark)
        
        # Log theme application
        logger.info(f"Applied {theme} theme to all widgets")
    
    def apply_theme_to_widget(self, widget, is_dark):
        """Recursively apply theme to a widget and all its children."""
        # Apply specific styling based on widget type
        if hasattr(widget, 'setProperty'):
            widget.setProperty("dark_mode", is_dark)
            
        # Process all child widgets
        for child in widget.findChildren(QWidget):
            self.apply_theme_to_widget(child, is_dark)
    
    def toggle_theme(self):
        """Toggle between light and dark themes."""
        # Toggle theme using theme manager
        theme, _ = self.theme_manager.toggle_theme()
        
        # Apply theme to all widgets
        self.apply_theme_to_all_widgets()
        
        # Save theme preference
        self.preferences["theme"] = theme
        
        # Update preferences
        self.preferences["theme"] = theme
        config.save_user_preferences(self.preferences)
        
        # Log theme change
        logger.info(f"Theme toggled to {theme} mode with background: {self.theme_manager.get_color('PrimarySolidBackgroundColor')}")
        
    # Removed update_suggestion_label_style method
    # as it is no longer needed with the new SuggestionPopup implementation
            
    # Removed update_suggestion_buttons_theme method
    # as it is no longer needed with the new SuggestionPopup implementation
        
    def update_icons_for_theme(self, theme):
        """Update all toolbar and menu icons for the current theme."""
        # Import the icon creation function
        from ui.icons import get_toolbar_icon
        
        # Update standard toolbar icons
        if hasattr(self, 'standard_toolbar'):
            for action in self.standard_toolbar.actions():
                icon_name = action.property("icon_name")
                if icon_name:
                    action.setIcon(get_toolbar_icon(icon_name, theme))
        
        # Update formatting toolbar icons
        if hasattr(self, 'formatting_toolbar'):
            for action in self.formatting_toolbar.actions():
                icon_name = action.property("icon_name")
                if icon_name:
                    action.setIcon(get_toolbar_icon(icon_name, theme))
        
        # Update toggles toolbar icons
        if hasattr(self, 'toggles_toolbar'):
            for action in self.toggles_toolbar.actions():
                icon_name = action.property("icon_name")
                if icon_name:
                    action.setIcon(get_toolbar_icon(icon_name, theme))
        
        # Update menu icons
        for menu in self.menuBar().findChildren(QMenu):
            for action in menu.actions():
                icon_name = action.property("icon_name")
                if icon_name:
                    action.setIcon(get_toolbar_icon(icon_name, theme))

    def update_recent_files_menu(self):
        """Update the recent files menu with the latest files."""
        # Create the menu if it doesn't exist
        if not self.recent_files_menu:
            return

        # Clear the menu
        self.recent_files_menu.clear()

        # Add recent files
        if self.preferences["recent_files"]:
            for filepath in self.preferences["recent_files"]:
                if os.path.exists(filepath):
                    action = QAction(os.path.basename(filepath), self)
                    action.setData(filepath)
                    action.triggered.connect(self.open_recent_file)
                    self.recent_files_menu.addAction(action)

            # Add separator and clear action
            if self.recent_files_menu.actions():
                self.recent_files_menu.addSeparator()
                clear_action = QAction("Clear Recent Files", self)
                clear_action.triggered.connect(self.clear_recent_files)
                self.recent_files_menu.addAction(clear_action)
        else:
            # Add a disabled "No Recent Files" action
            no_files_action = QAction("No Recent Files", self)
            no_files_action.setEnabled(False)
            self.recent_files_menu.addAction(no_files_action)

    def open_recent_file(self):
        """Open a file from the recent files menu."""
        action = self.sender()
        if action:
            filepath = action.data()
            if os.path.exists(filepath):
                # Check if there are unsaved changes
                if self.editor.document().isModified():
                    reply = QMessageBox.question(
                        self, 
                        "Save Changes?",
                        "Do you want to save changes before opening a new file?",
                        QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
                    )

                    if reply == QMessageBox.Cancel:
                        return
                    elif reply == QMessageBox.Yes:
                        self.save_file()

                # Open the file
                try:
                    # Determine file type based on extension
                    file_ext = os.path.splitext(filepath)[1].lower()
                    
                    # Read content based on file type
                    if file_ext == '.docx':
                        self._load_docx_into_editor(filepath)  # NEW
                        self.current_file = filepath
                        self.editor.document().setModified(False)
                        self.statusBar().showMessage("Loaded DOCX with styles", 2000)
                    elif file_ext == '.pdf':
                        content = read_pdf_file(filepath)
                        self.editor.setPlainText(content)
                    else:  # Default to text file
                        with open(filepath, "r", encoding="utf-8") as f:
                            self.editor.setPlainText(f.read())

                    # Update window title
                    self.setWindowTitle(f"Sinhala Word Processor - {os.path.basename(filepath)}")

                    # Move this file to the top of recent files
                    self.preferences = config.add_recent_file(self.preferences, filepath)
                    config.save_user_preferences(self.preferences)

                    # Update recent files menu
                    self.update_recent_files_menu()

                    logger.info(f"Opened recent file: {filepath}")
                except ImportError as e:
                    QMessageBox.critical(self, "Missing Dependency", f"{e}\nPlease restart the application to install required dependencies.")
                    logger.error(f"Dependency error opening file: {e}")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Could not open file:\n{e}")
                    logger.error(f"Error opening recent file: {e}")

    def clear_recent_files(self):
        """Clear the recent files list."""
        self.preferences["recent_files"] = []
        config.save_user_preferences(self.preferences)
        self.update_recent_files_menu()
        logger.info("Cleared recent files list")

    def build_menu(self):
        """Creates the application menu."""
        menu_bar = self.menuBar()

        # File Menu
        file_menu = menu_bar.addMenu("&File")
        # Use the actions stored as instance attributes
        file_menu.addAction(self.new_action)
        file_menu.addAction(self.open_action)
        file_menu.addAction(self.save_action)

        # Add Save As action
        self.save_as_action = QAction("Save As...", self)
        self.save_as_action.setShortcut("Ctrl+Shift+S")
        self.save_as_action.triggered.connect(self.save_as_file)
        file_menu.addAction(self.save_as_action)

        # Add Print and Print Preview actions
        file_menu.addAction(self.print_action)
        file_menu.addAction(self.print_preview_action)

        # Add Recent Files submenu
        file_menu.addSeparator()
        self.recent_files_menu = file_menu.addMenu("Recent Files")
        self.update_recent_files_menu()

        file_menu.addSeparator()
        # Create exit action here as it's only used in the menu
        exit_action = QAction("Exit", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # Edit Menu
        edit_menu = menu_bar.addMenu("&Edit")
        # Use the actions stored as instance attributes
        edit_menu.addAction(self.cut_action)
        edit_menu.addAction(self.copy_action)
        edit_menu.addAction(self.paste_action)
        edit_menu.addSeparator()
        edit_menu.addAction(self.undo_action)
        edit_menu.addAction(self.redo_action)

        # Add spell check action
        edit_menu.addSeparator()
        spell_check_action = QAction("Check Spelling", self)
        spell_check_action.setShortcut("F7")
        spell_check_action.triggered.connect(self.perform_spell_check)
        edit_menu.addAction(spell_check_action)

        # Add dictionary management action
        add_to_dict_action = QAction("Add Selected Word to Dictionary", self)
        add_to_dict_action.triggered.connect(self.learn_selected_word)
        edit_menu.addAction(add_to_dict_action)

        # View Menu
        view_menu = menu_bar.addMenu("&View")
        # Use the action stored as instance attribute
        view_menu.addAction(self.toggle_theme_action)

        # Add toggle toolbars action
        self.toggle_toolbars_action = QAction("Toggle Toolbars", self)
        self.toggle_toolbars_action.setShortcut("Ctrl+Alt+T")
        self.toggle_toolbars_action.triggered.connect(self.toggle_toolbars)
        view_menu.addAction(self.toggle_toolbars_action)

        # Add toggle keyboard action
        self.toggle_keyboard_action = QAction("Toggle Keyboard", self)
        self.toggle_keyboard_action.setShortcut("Ctrl+K")
        self.toggle_keyboard_action.triggered.connect(self.toggle_keyboard)
        view_menu.addAction(self.toggle_keyboard_action)
        
        # Add detach keyboard action
        self.detach_keyboard_action = QAction("Detach Keyboard", self)
        self.detach_keyboard_action.setShortcut("Ctrl+D")
        self.detach_keyboard_action.triggered.connect(self.detach_keyboard)
        view_menu.addAction(self.detach_keyboard_action)
        
        # Add dock keyboard action
        self.dock_keyboard_action = QAction("Dock Keyboard", self)
        self.dock_keyboard_action.setShortcut("Ctrl+Shift+D")
        self.dock_keyboard_action.triggered.connect(self.dock_keyboard)
        view_menu.addAction(self.dock_keyboard_action)
        
        # Add reset keyboard size action
        self.reset_keyboard_action = QAction("Set Keyboard (Size 20, Height 225px)", self)
        self.reset_keyboard_action.setShortcut("Ctrl+Alt+K")
        self.reset_keyboard_action.triggered.connect(self.reset_keyboard_size)
        view_menu.addAction(self.reset_keyboard_action)
        
        # Add keyboard font size submenu
        keyboard_font_menu = QMenu("Keyboard Font Size", self)
        
        # Add font size options based on constants
        # Start with the default size
        default_action = QAction(f"{DEFAULT_KB_FONT_SIZE}pt (Default)", self)
        default_action.triggered.connect(lambda checked: self.set_keyboard_font_size(DEFAULT_KB_FONT_SIZE))
        keyboard_font_menu.addAction(default_action)
        
        # Add separator
        keyboard_font_menu.addSeparator()
        
        # Add other size options
        for size in [MIN_KB_FONT, 30, 35, 40, 45, 50, 60, 70, 80]:
            # Skip the default size as it's already added
            if size == DEFAULT_KB_FONT_SIZE:
                continue
                
            action = QAction(f"{size}pt", self)
            action.triggered.connect(lambda checked, s=size: self.set_keyboard_font_size(s))
            keyboard_font_menu.addAction(action)
            
        view_menu.addMenu(keyboard_font_menu)

        # Add toggle suggestions action
        view_menu.addAction(self.suggestions_toggle_action)

        # Add toggle Singlish action
        view_menu.addAction(self.singlish_toggle_action)
        
        # Add separator
        view_menu.addSeparator()
        
        # Add settings action
        self.settings_action = QAction("Settings...", self)
        self.settings_action.setShortcut("Ctrl+,")
        self.settings_action.triggered.connect(self.show_settings_dialog)
        view_menu.addAction(self.settings_action)

        # Help Menu
        help_menu = menu_bar.addMenu("&Help")

        # About action
        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about_dialog)
        help_menu.addAction(about_action)

        # Help action
        help_action = QAction("Help", self)
        help_action.setShortcut("F1")
        help_action.triggered.connect(self.show_help_dialog)
        help_menu.addAction(help_action)

    def show_about_dialog(self):
        """Show the about dialog."""
        QMessageBox.about(
            self,
            "About Sinhala Word Processor",
            """<h1>Sinhala Word Processor</h1>
            <p>Version 3.3</p>
            <p>A lightweight word processor with Sinhala transliteration support.</p>
            <p>Features:</p>
            <ul>
                <li>Realtime Singlish to Sinhala transliteration</li>
                <li>Suggestion popup</li>
                <li>Basic spell checker</li>
                <li>Dark/Light theme toggle</li>
            </ul>
            """
        )

    def show_help_dialog(self):
        """Show the help dialog."""
        QMessageBox.information(
            self,
            "Help",
            """<h2>Sinhala Word Processor Help</h2>
            <h3>Keyboard Shortcuts:</h3>
            <ul>
                <li><b>Ctrl+N</b>: New file</li>
                <li><b>Ctrl+O</b>: Open file</li>
                <li><b>Ctrl+S</b>: Save file</li>
<li><b>Ctrl+Shift+S</b>: Save file as</li><li><b>Ctrl+X</b>: Cut</li>
                <li><b>Ctrl+C</b>: Copy</li>
                <li><b>Ctrl+V</b>: Paste</li>
                <li><b>Ctrl+Z</b>: Undo</li>
                <li><b>Ctrl+Y</b>: Redo</li>
                <li><b>Ctrl+T</b>: Toggle theme</li>
                <li><b>Ctrl+K</b>: Toggle keyboard</li>
                <li><b>F7</b>: Check spelling</li>
            </ul>

            <h3>Singlish to Sinhala:</h3>
            <p>Type in Singlish (romanized Sinhala) and it will be automatically 
            converted to Sinhala. Use numbers 1-9 to select from suggestions.</p>

            <h3>Adding Words to Dictionary:</h3>
            <p>Select a Sinhala word, right-click and choose "Learn Selected Word" 
            to add it to your personal dictionary.</p>
            """
        )

    def new_file(self):
        """Create a new file."""
        # Ask to save changes if there's content in the editor
        if self.editor.document().isModified():
            reply = QMessageBox.question(
                self, 
                "Save Changes?",
                "Do you want to save changes before creating a new file?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
            )

            if reply == QMessageBox.Cancel:
                return
            elif reply == QMessageBox.Yes:
                self.save_file()

        # Clear the editor
        self.editor.clear()
        self.buffer.clear()
        self.word_start_pos = None
        self.clear_suggestion_area()
        self.update_status()

        # Reset window title
        self.setWindowTitle("Sinhala Word Processor")

    # ---------- DOCX reader that keeps styles ------------------------------
    def _load_docx_into_editor(self, path: str):
        """Clear the editor and re-create the DOCX with fonts / bold / italic / underline."""
        docx = Document(path)
        self.editor.clear()
        cursor: QTextCursor = self.editor.textCursor()
        for para in docx.paragraphs:
            for run in para.runs:
                fmt = QTextCharFormat()
                qf = QFont()
                # family / size
                if run.font.name:
                    qf.setFamily(run.font.name)
                if run.font.size:
                    qf.setPointSizeF(run.font.size.pt)  # python-docx size → pt
                # inline styles
                # --- replace the three style setters ------------------------
                qf.setBold(bool(run.bold))  # was: qf.setBold(run.bold)
                qf.setItalic(bool(run.italic))  # was: qf.setItalic(run.italic)
                qf.setUnderline(bool(run.underline))  # was: qf.setUnderline(run.underline)
                # ------------------------------------------------------------
                fmt.setFont(qf)
                fmt.setFontUnderline(bool(run.underline))  # ← ADD: explicit underline flag on QTextCharFormat
                cursor.setCharFormat(fmt)
                cursor.insertText(run.text)
            cursor.insertBlock()  # paragraph break
    
    def open_file(self):
        """Opens a file in the editor with support for multiple formats (txt, docx, pdf)."""
        # Ask to save changes if there's content in the editor
        if self.editor.document().isModified():
            reply = QMessageBox.question(
                self, 
                "Save Changes?",
                "Do you want to save changes before opening a new file?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
            )

            if reply == QMessageBox.Cancel:
                return
            elif reply == QMessageBox.Yes:
                self.save_file()

        # Get the directory of the most recently opened file, if any
        initial_dir = ""
        if self.preferences["recent_files"] and os.path.exists(os.path.dirname(self.preferences["recent_files"][0])):
            initial_dir = os.path.dirname(self.preferences["recent_files"][0])

        file_path, selected_filter = QFileDialog.getOpenFileName(
            self, 
            "Open File", 
            initial_dir, 
            "All Supported Files (*.txt *.docx *.pdf);;Text Files (*.txt);;Word Documents (*.docx);;PDF Files (*.pdf);;All Files (*)"
        )

        if file_path:
            try:
                # Determine file type based on extension
                file_ext = os.path.splitext(file_path)[1].lower()
                
                # Read content based on file type
                if file_ext == '.docx':
                    self._load_docx_into_editor(file_path)  # NEW
                    self.current_file = file_path
                    self.editor.document().setModified(False)
                    self.statusBar().showMessage("Loaded DOCX with styles", 2000)
                    return
                elif file_ext == '.pdf':
                    content = read_pdf_file(file_path)
                else:  # Default to text file
                    content = read_text_file(file_path)
                
                # Set the content in the editor
                self.editor.setPlainText(content)

                # Add to recent files
                self.preferences = config.add_recent_file(self.preferences, file_path)
                config.save_user_preferences(self.preferences)

                # Update recent files menu
                self.update_recent_files_menu()

                # Update window title with filename
                self.setWindowTitle(f"Sinhala Word Processor - {os.path.basename(file_path)}")

                logger.info(f"Opened file: {os.path.basename(file_path)}")

                # Perform spell check on the loaded file
                self.spell_check_timer.start(500)  # Start spell check after 500ms
            except ImportError as e:
                QMessageBox.critical(self, "Missing Dependency", f"{e}\nPlease restart the application to install required dependencies.")
                logger.error(f"Dependency error opening file: {e}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not open file:\n{e}")
                logger.error(f"Error opening file: {e}")

    def save_file(self):
        """Saves the current editor content to a file with support for multiple formats (txt, docx, pdf)."""
        # Check if we have a current file path from recent files
        current_file = None
        if self.windowTitle() != "Sinhala Word Processor":
            # Extract filename from window title
            filename = self.windowTitle().replace("Sinhala Word Processor - ", "")
            # Find the file in recent files
            for filepath in self.preferences["recent_files"]:
                if os.path.basename(filepath) == filename:
                    current_file = filepath
                    break

        if current_file and os.path.exists(current_file):
            # Save to the current file
            try:
                file_ext = os.path.splitext(current_file)[1].lower()
                
                # Use the new _do_save method
                try:
                    self._do_save(current_file)
                except Exception as err:
                    QMessageBox.critical(self, "Save error", str(err))
                    return False
                
                logger.info(f"Saved file: {os.path.basename(current_file)}")
                self.editor.document().setModified(False)
                return True
            except ImportError as e:
                QMessageBox.critical(self, "Missing Dependency", f"{e}\nPlease restart the application to install required dependencies.")
                logger.error(f"Dependency error saving file: {e}")
                return False
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not save file:\n{e}")
                logger.error(f"Error saving file: {e}")
                return False
        else:
            # No current file, use Save As
            return self.save_as_file()

    def save_as_file(self):
        """Save the current editor content to a new file with support for multiple formats (txt, docx, pdf)."""
        # Get the directory of the most recently saved file, if any
        initial_dir = ""
        if self.preferences["recent_files"] and os.path.exists(os.path.dirname(self.preferences["recent_files"][0])):
            initial_dir = os.path.dirname(self.preferences["recent_files"][0])

        file_path, selected_filter = QFileDialog.getSaveFileName(
            self, 
            "Save File As", 
            initial_dir, 
            "All Supported Files (*.txt *.docx *.pdf);;Text Files (*.txt);;Word Documents (*.docx);;PDF Files (*.pdf);;All Files (*)"
        )

        if file_path:
            try:
                content = self.editor.toPlainText()
                
                # Determine file type based on extension
                file_ext = os.path.splitext(file_path)[1].lower()
                
                # If no extension was provided, add one based on the selected filter
                if not file_ext:
                    if "Word Documents" in selected_filter:
                        file_path += ".docx"
                        file_ext = ".docx"
                    elif "PDF Files" in selected_filter:
                        file_path += ".pdf"
                        file_ext = ".pdf"
                    elif "Text Files" in selected_filter:
                        file_path += ".txt"
                        file_ext = ".txt"
                
                # Use the new _do_save method
                try:
                    self._do_save(file_path)
                except Exception as err:
                    QMessageBox.critical(self, "Save error", str(err))
                    return False

                # Add to recent files
                self.preferences = config.add_recent_file(self.preferences, file_path)
                config.save_user_preferences(self.preferences)

                # Update recent files menu
                self.update_recent_files_menu()

                # Update window title with filename
                self.setWindowTitle(f"Sinhala Word Processor - {os.path.basename(file_path)}")

                logger.info(f"Saved file as: {os.path.basename(file_path)}")
                self.editor.document().setModified(False)
                return True
            except ImportError as e:
                QMessageBox.critical(self, "Missing Dependency", f"{e}\nPlease restart the application to install required dependencies.")
                logger.error(f"Dependency error saving file: {e}")
                return False
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not save file:\n{e}")
                logger.error(f"Error saving file: {e}")
                return False

        return False

    def _save_docx(self, path: str):
        try:
            qdoc = self.editor.document()
            docx = Document()
            block = qdoc.begin()
            
            while block != qdoc.end():
                p = docx.add_paragraph()
                it = block.begin()
                
                while not it.atEnd():
                    frag = it.fragment()
                    if not frag.isValid():
                        it += 1
                        continue
                        
                    text = frag.text()
                    if text == "":
                        it += 1
                        continue
                        
                    fmt = frag.charFormat()
                    qf = fmt.font()
                    run = p.add_run(text)
                    
                    # font family / size
                    run.font.name = qf.family() or self.base_font.family()
                    run.font.size = Pt(qf.pointSizeF() or self.base_font.pointSizeF())
                    
                    # inline styles
                    run.bold = qf.bold()
                    run.italic = qf.italic()
                    run.underline = fmt.fontUnderline()  # replace run.underline = qf.underline()
                    
                    # colour
                    col = fmt.foreground().color()
                    if col.isValid():
                        run.font.color.rgb = RGBColor(col.red(), col.green(), col.blue())
                        
                    it += 1
                    
                block = block.next()
                
            docx.save(path)  # <-- if this raises, we catch below
            
        except Exception as e:
            QMessageBox.critical(self, "DOCX export failed", str(e))
            raise  # Re-raise to be caught by the outer try-except

    def _do_save(self, path: str):
        ext = Path(path).suffix.lower()
        if ext == ".docx":
            self._save_docx(path)  # call the new safe writer
        elif ext == ".pdf":
            document = self.editor.document()
            write_pdf_file(path, document)
        else:  # Default to text file
            content = self.editor.toPlainText()
            write_text_file(path, content)

    def handle_keypress_event(self, obj, event):
        """Handle keypress events for the editor."""
        try:
            key = event.key()
            text = event.text()

            # Check if Singlish is enabled
            singlish_enabled = hasattr(self, 'singlish_toggle_action') and self.singlish_toggle_action.isChecked()
            suggestions_enabled = hasattr(self, 'suggestions_toggle_action') and self.suggestions_toggle_action.isChecked()

            # If Singlish is disabled, don't process any special handling
            if not singlish_enabled:
                return False
                
            # Handle cursor movement keys - commit buffer if active
            if key in (Qt.Key_Left, Qt.Key_Right, Qt.Key_Up, Qt.Key_Down, Qt.Key_Home, Qt.Key_End) and self.buffer:
                logger.info(f"Cursor movement detected with active buffer - committing buffer")
                self.commit_buffer()
                return False  # Don't consume the event
                
            # If suggestions are enabled and we have suggestions, handle selection keys
            if suggestions_enabled and self.current_suggestions:
                if key in (Qt.Key_Return, Qt.Key_Enter):
                    # Accept the current suggestion on Enter
                    try:
                        # Get the current index from the popup
                        current_index = self.suggestion_popup.current_index
                        if current_index >= 0 and current_index < len(self.current_suggestions):
                            sinhala_word = self.current_suggestions[current_index]
                        else:
                            # Default to first suggestion if no current index
                            sinhala_word = self.current_suggestions[0]
                        
                        # Log the selected suggestion
                        logger.info(f"Selected suggestion: '{sinhala_word}'")
                        
                        # Clear suggestions first to avoid race conditions
                        self.clear_suggestion_area()
                        # Accept the suggestion (this will also clear the buffer)
                        self.accept_suggestion(sinhala_word)
                        return True # Consume Enter
                    except Exception as e:
                        logger.error(f"Error handling Enter key: {e}")
                        self.reset_input_state()
                        return False
                        
                elif key == Qt.Key_Space:
                    try:
                        # Check if the buffer starts with a number
                        buffer_text = "".join(self.buffer) if self.buffer else ""
                        is_numeric_input = buffer_text and buffer_text[0].isdigit()
                        
                        # For numeric input, don't use the suggestion system
                        if is_numeric_input:
                            logger.info(f"Space pressed after numeric input: '{buffer_text}' - treating as regular input")
                            # Just clear the buffer without committing
                            self.buffer.clear()
                            self.word_start_pos = None
                            self.clear_suggestion_area()
                            # Let the editor handle the space normally
                            return False
                        
                        # When space is pressed with suggestions visible
                        if self.suggestion_popup.isVisible() and self.current_suggestions:
                            # Hide the suggestion popup
                            self.suggestion_popup.hide()
                            # Clear suggestions
                            self.current_suggestions = []
                            
                            # Safely commit the buffer if it exists
                            if self.buffer:
                                # Make sure we have a valid word_start_pos before committing
                                if self.word_start_pos is not None:
                                    self.commit_buffer()
                                else:
                                    # If word_start_pos is None, just clear the buffer
                                    logger.warning("Space pressed with buffer but no word_start_pos - clearing buffer")
                                    self.buffer.clear()
                            
                            # Insert a space
                            self.editor.insertPlainText(" ")
                            return True # Consume Space
                        else:
                            # No visible suggestions, commit buffer if any and let space through
                            if self.buffer:
                                # Make sure we have a valid word_start_pos before committing
                                if self.word_start_pos is not None:
                                    self.commit_buffer()
                                else:
                                    # If word_start_pos is None, just clear the buffer
                                    logger.warning("Space pressed with buffer but no word_start_pos - clearing buffer")
                                    self.buffer.clear()
                            return False
                    except Exception as e:
                        logger.error(f"Error handling Space key: {e}")
                        import traceback
                        logger.error(f"Traceback: {traceback.format_exc()}")
                        self.reset_input_state()
                        return False
                        
                elif key == Qt.Key_Escape:
                    self.reset_input_state()
                    return True # Consume Escape
                    
                elif key == Qt.Key_Tab:
                    # Navigate to next suggestion
                    self.suggestion_popup.navigate_next()
                    return True # Consume Tab
                    
                elif key == Qt.Key_Backtab:
                    # Navigate to previous suggestion
                    self.suggestion_popup.navigate_previous()
                    return True # Consume Shift+Tab
                    
                elif key == Qt.Key_Down:
                    # Navigate to next suggestion
                    self.suggestion_popup.navigate_next()
                    return True # Consume Down arrow
                    
                elif key == Qt.Key_Up:
                    # Navigate to previous suggestion
                    self.suggestion_popup.navigate_previous()
                    return True # Consume Up arrow
                    
                elif Qt.Key_1 <= key <= Qt.Key_9:
                    # Only process number keys for suggestion selection when the popup is visible
                    if self.suggestion_popup.isVisible():
                        index = key - Qt.Key_1
                        if index < len(self.current_suggestions):
                            try:
                                # Make a local copy of the suggestion to avoid race conditions
                                sinhala_word = self.current_suggestions[index]
                                
                                # Log the selected suggestion
                                logger.info(f"Selected suggestion {index+1}: '{sinhala_word}'")
                                
                                # Clear suggestions first to avoid race conditions
                                self.clear_suggestion_area()
                                # Accept the suggestion (this will also clear the buffer)
                                self.accept_suggestion(sinhala_word)
                                return True # Consume number key
                            except Exception as e:
                                logger.error(f"Error handling number key: {e}")
                                logger.error(f"Exception details: {str(e)}")
                                import traceback
                                logger.error(f"Traceback: {traceback.format_exc()}")
                                self.reset_input_state()
                                return False
                    else:
                        # If popup is not visible, treat number keys as regular input
                        # Let the alphanumeric handler below process it
                        logger.info(f"Number key pressed but suggestion popup not visible - treating as regular input")
                        return False

            # If Backspace is pressed, handle buffer modification
            if key == Qt.Key_Backspace:
                # If buffer is not empty, remove last char from buffer
                if self.buffer:
                    self.buffer.pop()
                    # After popping, update or clear suggestion area
                    if suggestions_enabled:
                        if self.buffer:
                            # Cancel any pending suggestion updates
                            if hasattr(self, '_suggestion_timer') and self._suggestion_timer.isActive():
                                self._suggestion_timer.stop()
                            
                            # Create a new timer with a slightly longer delay
                            self._suggestion_timer = QTimer()
                            self._suggestion_timer.setSingleShot(True)
                            self._suggestion_timer.timeout.connect(self.update_suggestion_area)
                            self._suggestion_timer.start(30)  # Increased delay for better stability
                        else:
                            self.clear_suggestion_area()
                            self.word_start_pos = None
                    # Do NOT consume backspace here, let the editor handle it
                    return False
                else:
                    # If buffer is empty, let editor handle backspace normally
                    if suggestions_enabled:
                        self.clear_suggestion_area() # Ensure area is hidden if buffer becomes empty
                    self.word_start_pos = None # Explicitly reset word_start_pos
                    return False

            # Handle alphanumeric keys: append to buffer
            if text.isalnum() and len(text) == 1:
                # Get current cursor position before any operations
                current_pos = self.editor.textCursor().position()
                
                # Check if this is a numeric character starting a new buffer
                is_numeric = text.isdigit()
                is_new_buffer = not self.buffer
                
                # If this is a numeric character starting a new buffer, handle it differently
                if is_numeric and is_new_buffer:
                    # For numeric input at the start of a buffer, don't use suggestions
                    logger.info(f"Numeric character at start of buffer: '{text}' - treating as regular input")
                    # Let the editor handle it normally, don't add to buffer
                    return False
                
                # Verify cursor position is where we expect it to be if buffer is not empty
                if self.buffer and self.word_start_pos is not None:
                    expected_pos = self.word_start_pos + len(self.buffer)
                    
                    # If cursor is not where expected, commit buffer and start a new one
                    if current_pos != expected_pos:
                        logger.info(f"Cursor position mismatch: expected {expected_pos}, got {current_pos} - committing buffer")
                        self.commit_buffer()
                        # Set word_start_pos to current position AFTER committing buffer
                        self.word_start_pos = current_pos
                
                # If buffer was empty, mark the start position
                # This needs to happen AFTER any potential buffer commit
                if not self.buffer:
                    self.word_start_pos = current_pos
                    logger.info(f"Starting new buffer at position {current_pos}")
                     
                # Now append to buffer
                self.buffer.append(text)
                
                # Check if the buffer now starts with a number
                buffer_text = "".join(self.buffer)
                if buffer_text and buffer_text[0].isdigit():
                    # For buffers starting with numbers, don't show suggestions
                    if suggestions_enabled:
                        self.clear_suggestion_area()
                else:
                    # Update suggestions immediately after adding to buffer (only for non-numeric input)
                    if suggestions_enabled:
                        # Cancel any pending suggestion updates
                        if hasattr(self, '_suggestion_timer') and self._suggestion_timer.isActive():
                            self._suggestion_timer.stop()
                        
                        # Use a very short timer to ensure cursor position is stable
                        self._suggestion_timer = QTimer()
                        self._suggestion_timer.setSingleShot(True)
                        self._suggestion_timer.timeout.connect(self.update_suggestion_area)
                        self._suggestion_timer.start(10)  # Very short delay to avoid race conditions
                
                # Do NOT consume the event, let the editor insert the character
                return False

            # For any other key (punctuation, symbols, arrows, etc.), commit the current buffer
            # and then clear the suggestion area.
            if self.buffer:
                self.commit_buffer()
            return False # Do not consume the event, let the editor handle it

        except Exception as e:
            logger.error(f"Exception in handle_keypress_event: {e}")
            # Ensure area is hidden on error
            self.reset_input_state()
            return False
            
    def reset_input_state(self):
        """Reset all input state variables to avoid crashes."""
        # Cancel any pending suggestion updates
        if hasattr(self, '_suggestion_timer') and self._suggestion_timer.isActive():
            self._suggestion_timer.stop()
            
        self.buffer.clear()
        self.word_start_pos = None
        self.clear_suggestion_area()
        
        logger.info("Input state reset")
        
    def replace_with_suggestion(self, singlish_word):
        """Replace the current buffer text with the selected suggestion."""
        try:
            if not singlish_word:
                return
                
            # Get the Sinhala word for the suggestion
            sinhala_word = self.transliterator.transliterate(singlish_word)
            if not sinhala_word or sinhala_word == singlish_word:
                # Fallback to phonetic if not in map
                sinhala_word = _phonetic_global(singlish_word)
                
            if not sinhala_word:
                return
                
            # Only proceed if we have a valid word_start_pos
            if self.word_start_pos is None:
                return
                
            # Get the current cursor
            cur = self.editor.textCursor()
            cur.beginEditBlock()
            
            # Store the current position
            current_pos = cur.position()
            
            # Get the current document text
            doc_text = self.editor.toPlainText()
            
            # Calculate how many characters to select - use the buffer content length
            # instead of the suggestion length, as the buffer contains what was actually typed
            buffer_text = "".join(self.buffer) if self.buffer else singlish_word
            buffer_length = len(buffer_text)
            
            # Make sure we don't go beyond document bounds
            document_length = len(doc_text)
            if self.word_start_pos >= document_length:
                # Invalid position, abort
                cur.endEditBlock()
                return
                
            # Adjust buffer_length to not exceed document bounds
            if self.word_start_pos + buffer_length > document_length:
                buffer_length = document_length - self.word_start_pos
                
            # Move to the start position
            cur.setPosition(self.word_start_pos)
            
            # Select the text to replace (only if we have a valid length)
            if buffer_length > 0:
                cur.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, buffer_length)
                
                # Replace with the Sinhala word
                cur.insertText(sinhala_word)
            
            cur.endEditBlock()
            
        except Exception as e:
            logger.error(f"Error in replace_with_suggestion: {e}")
            # Reset state on error
            self.reset_input_state()

    def update_suggestion_area(self):
        """Update and show/hide the suggestion popup based on the current buffer."""
        try:
            # Check if suggestions are enabled
            if hasattr(self, 'suggestions_toggle_action') and not self.suggestions_toggle_action.isChecked():
                logger.info("Suggestions disabled")
                self.suggestion_popup.hide()
                return
            
            # If buffer is empty or word_start_pos is None, hide suggestion popup
            if not self.buffer or self.word_start_pos is None:
                logger.info("Buffer is empty or word_start_pos is None, hiding suggestion popup")
                self.suggestion_popup.hide()
                return
            
            # Get the current word from buffer
            current_word = "".join(self.buffer).lower()
            
            # Skip suggestion logic for numeric input
            if current_word and current_word[0].isdigit():
                logger.info(f"Numeric input detected: '{current_word}' - skipping suggestions")
                self.suggestion_popup.hide()
                self.current_suggestions = []  # Clear any existing suggestions
                return
            
            # Verify cursor position is still valid for this buffer
            current_pos = self.editor.textCursor().position()
            expected_pos = self.word_start_pos + len(self.buffer)
            
            # If cursor has moved away from where we expect it, don't update suggestions
            if current_pos != expected_pos:
                logger.info(f"Cursor position mismatch in update_suggestion_area: expected {expected_pos}, got {current_pos}")
                # Don't update suggestions if cursor is not where expected
                return
            
            # Use the transliterator to get suggestions
            self.current_suggestions = self.transliterator.get_suggestions(current_word, max_suggestions=9)
            
            # Log the suggestions for debugging
            logger.info(f"Suggestions for '{current_word}': {self.current_suggestions}")
            
            # Get the cursor rectangle to position the popup
            cursor = self.editor.textCursor()
            cursor_rect = self.editor.cursorRect(cursor)
            editor_rect = self.editor.rect()
            
            # Update theme colors for the popup
            theme_colors = {
                "OverlayColor": self.theme_manager.get_color("OverlayColor"),
                "PrimarySolidBackgroundColor": self.theme_manager.get_color("PrimarySolidBackgroundColor"),
                "PrimarySolidBorderColor": self.theme_manager.get_color("PrimarySolidBorderColor"),
                "PrimaryForegroundColor": self.theme_manager.get_color("PrimaryForegroundColor"),
                "MouseOverBackgroundColor": self.theme_manager.get_color("MouseOverBackgroundColor"),
                "AccentColor": self.theme_manager.get_color("AccentColor"),
                "AccentPressedColor": self.theme_manager.get_color("AccentPressedColor")
            }
            
            # Update the popup theme
            self.suggestion_popup.update_theme(
                self.theme_manager.is_dark_mode(),
                theme_colors
            )
            
            # Show the popup with suggestions
            if self.current_suggestions:
                # Save current cursor position
                saved_position = cursor.position()
                
                # Show the popup with suggestions
                self.suggestion_popup.show_popup(
                    self.current_suggestions,
                    cursor_rect,
                    editor_rect
                )
                
                # Make sure the editor still has focus after showing popup
                # But don't change the cursor position
                self.editor.setFocus()
                
                logger.info(f"Showing suggestion popup with {len(self.current_suggestions)} suggestions")
            else:
                # No suggestions available, hide the popup
                self.suggestion_popup.hide()
                logger.info("No suggestions available, hiding popup")
            
        except Exception as e:
            logger.error(f"Error in update_suggestion_area: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            # Ensure popup is hidden on error
            self.suggestion_popup.hide()
            self.current_suggestions = []

    def clear_suggestion_area(self):
        """Hide the suggestion popup and clear stored suggestions."""
        try:
            # Hide the popup
            self.suggestion_popup.hide()
            
            # Clear stored suggestions
            self.current_suggestions = []
            
            # Log that we've cleared the area
            logger.info("Cleared suggestion area")
        except Exception as e:
            logger.error(f"Error in clear_suggestion_area: {e}")
            # Reset everything to be safe
            self.current_suggestions = []
        
        # Log the clearing
        logger.info("Cleared suggestion area")

    def accept_suggestion(self, sinhala_word):
        """Accept a suggestion and replace the buffered text."""
        # This matches the behavior in SinhalaWordProcessor_enhanced.py
        if self.buffer and self.word_start_pos is not None:
            try:
                # Log the operation
                buffer_text = "".join(self.buffer)
                logger.info(f"Accepting suggestion: '{sinhala_word}' to replace buffer: '{buffer_text}'")
                
                cur = self.editor.textCursor()
                cur.beginEditBlock()

                # Get the actual text in the document at the expected position
                doc = self.editor.document()
                check_cursor = QTextCursor(doc)
                check_cursor.setPosition(self.word_start_pos)
                check_cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(buffer_text))
                actual_text = check_cursor.selectedText()
                
                # Verify the text matches our buffer before replacing
                if actual_text == buffer_text:
                    # Standard replacement as before
                    cur.setPosition(self.word_start_pos)
                    cur.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(buffer_text))
                    cur.removeSelectedText()
                    cur.insertText(sinhala_word)
                    logger.info(f"Text matched buffer - standard replacement performed")
                else:
                    # Text doesn't match buffer - use a more robust approach
                    logger.warning(f"Buffer text '{buffer_text}' doesn't match document text '{actual_text}'")
                    
                    # Try to find the buffer text near the current position
                    current_pos = cur.position()
                    search_start = max(0, self.word_start_pos - 10)
                    search_end = min(doc.characterCount(), current_pos + 10)
                    
                    # Create a cursor for searching
                    search_cursor = QTextCursor(doc)
                    search_cursor.setPosition(search_start)
                    search_cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, search_end - search_start)
                    search_text = search_cursor.selectedText()
                    
                    # Try to find the buffer text in the search area
                    buffer_pos = search_text.find(buffer_text)
                    if buffer_pos >= 0:
                        # Found the buffer text, replace it
                        replace_pos = search_start + buffer_pos
                        cur.setPosition(replace_pos)
                        cur.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(buffer_text))
                        cur.removeSelectedText()
                        cur.insertText(sinhala_word)
                        logger.info(f"Found buffer text at position {replace_pos} - replacement performed")
                    else:
                        # Couldn't find the buffer text, just insert at current position
                        logger.warning(f"Couldn't find buffer text near cursor - inserting at current position")
                        cur.insertText(sinhala_word)
                
                cur.endEditBlock()

                # Clear the buffer and reset word_start_pos
                self.buffer.clear()
                self.word_start_pos = None
                
                # Log success
                logger.info(f"Successfully accepted suggestion: '{sinhala_word}'")
            except Exception as e:
                logger.error(f"Error in accept_suggestion: {e}")
                logger.error(f"Exception details: {str(e)}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
                self.reset_input_state()
                
                # Clear area after accepting
                self.clear_suggestion_area()
        else:
            # If there's no buffer or word_start_pos, just clear the suggestion area
            self.clear_suggestion_area()

    def commit_buffer(self):
        """Commit the buffered input to the editor (phonetic transliteration)."""
        try:
            # Check if Singlish is enabled
            singlish_enabled = hasattr(self, 'singlish_toggle_action') and self.singlish_toggle_action.isChecked()
            
            # Clear suggestion area regardless of Singlish state
            self.clear_suggestion_area()
            
            # If Singlish is disabled or buffer is empty, don't process
            if not singlish_enabled or not self.buffer or self.word_start_pos is None:
                # Always clear the buffer and reset word_start_pos for safety
                self.buffer.clear()
                self.word_start_pos = None
                return

            word = "".join(self.buffer)
            logger.info(f"Committing buffer: '{word}' at position {self.word_start_pos}")

            # Get the document for validation
            doc = self.editor.document()
            
            # Validate word_start_pos before proceeding
            if not isinstance(self.word_start_pos, int) or self.word_start_pos < 0 or self.word_start_pos >= doc.characterCount():
                logger.warning(f"Invalid word_start_pos: {self.word_start_pos} (document length: {doc.characterCount()}) - aborting commit")
                self.buffer.clear()
                self.word_start_pos = None
                return
                
            # Validate that the selection range is valid
            if self.word_start_pos + len(word) > doc.characterCount():
                logger.warning(f"Invalid selection range: start={self.word_start_pos}, length={len(word)}, doc length={doc.characterCount()} - aborting commit")
                self.buffer.clear()
                self.word_start_pos = None
                return

            # Use the transliterator to get the Sinhala word
            sinhala_word = self.transliterator.transliterate(word)
            if not sinhala_word or sinhala_word == word:
                # Fallback to phonetic
                sinhala_word = _phonetic_global(word)
                
            logger.info(f"Transliterated to: '{sinhala_word}'")

            # Create a cursor for verification
            check_cursor = QTextCursor(doc)
            check_cursor.setPosition(self.word_start_pos)
            
            # Try to select the text that should match our buffer
            check_cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(word))
            actual_text = check_cursor.selectedText()
            
            # Verify the text matches our buffer before replacing
            if actual_text == word:
                # Text matches, do standard replacement
                try:
                    cur = self.editor.textCursor()
                    cur.beginEditBlock()
                    
                    # Save the current position to restore it later if needed
                    original_position = cur.position()
                    
                    # Move cursor to the start of the buffered word
                    cur.setPosition(self.word_start_pos)
                    # Select the buffered text
                    cur.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(word))
                    # Remove the buffered text
                    cur.removeSelectedText()
                    # Insert the transliterated Sinhala word
                    cur.insertText(sinhala_word)
                    cur.endEditBlock()
                    
                    logger.info(f"Successfully replaced '{word}' with '{sinhala_word}'")
                except Exception as inner_e:
                    logger.error(f"Error during text replacement: {inner_e}")
                    # Make sure we end the edit block even if there's an error
                    if cur.isNull() == False:
                        cur.endEditBlock()
            else:
                # Text doesn't match buffer - log the mismatch
                logger.warning(f"Buffer text '{word}' doesn't match document text '{actual_text}' - skipping replacement")

            # Always clear the buffer and reset word_start_pos
            self.buffer.clear()
            self.word_start_pos = None # Reset word_start_pos after committing
            
        except Exception as e:
            logger.error(f"Error in commit_buffer: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            # Reset state on error
            self.buffer.clear()
            self.word_start_pos = None
            self.clear_suggestion_area()

    def eventFilter(self, obj, event):
        # Check if the event is for the editor
        if event.type() == QEvent.KeyPress and (obj is self.editor or obj is self.editor.viewport()):
             # Pass the event and the object it occurred on to handle_keypress_event
             return self.handle_keypress_event(obj, event)
        
        # Track cursor position changes when buffer is active
        elif event.type() == QEvent.MouseButtonPress and (obj is self.editor or obj is self.editor.viewport()) and self.buffer:
            # If user clicks elsewhere while buffer is active, commit the buffer
            logger.info("Mouse click detected with active buffer - committing buffer")
            self.commit_buffer()
            return False  # Don't consume the event
            
        # For any other event or object, let the default event filter handle it
        return super().eventFilter(obj, event)

    @Slot() # Decorator to indicate this is a slot
    def on_text_changed(self):
        # The buffer is now managed by handle_keypress_event
        # This slot is mainly for updating the status bar
        self.update_status()

        # Start spell check timer
        self.spell_check_timer.start(1000)  # Check spelling after 1 second of inactivity

    def perform_spell_check(self):
        """Perform spell checking on the editor content."""
        # Get the current text
        text = self.editor.toPlainText()

        # --- We'll skip clearing underlines entirely ---
        # This way, we won't interfere with any manual formatting
        # The spell checker will just add red underlines to misspelled words
        
        # Store the current cursor position and selection
        old_cursor = self.editor.textCursor()
        old_position = old_cursor.position()
        old_anchor = old_cursor.anchor()
        
        # We'll let the spell checker add its underlines without clearing anything

        # Find all Sinhala words and check them
        sinhala_word_pattern = re.compile(r'[\u0D80-\u0DFF]+')
        for match in sinhala_word_pattern.finditer(text):
            word = match.group(0)
            if not self.spellchecker.is_known_word(word):
                # Underline the misspelled word
                cursor = self.editor.textCursor()
                cursor.setPosition(match.start())
                cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, len(word))

                error_fmt = QTextCharFormat()
                error_fmt.setUnderlineStyle(QTextCharFormat.SpellCheckUnderline)
                error_fmt.setUnderlineColor(QColor("red"))
                
                # add red underline to misspelled range
                cursor.beginEditBlock()
                cursor.mergeCharFormat(error_fmt)     # keep as merge, not set
                cursor.endEditBlock()

    def suggestions(self, prefix: str, limit: int = 9):
        """Get a list of suggestions for a given prefix from the dictionary."""
        # Use the transliterator to get suggestions
        return self.transliterator.get_suggestions(prefix, max_suggestions=limit)

    # --- Transliteration/Suggestion Methods ---
    def tr(self, word: str) -> str:
        """Transliterate a word from Singlish to Sinhala."""
        return self.transliterator.transliterate(word) or _phonetic_global(word)

    # --- Save Method ---
    def _save_map(self, force=False):
        if force or self.SAVE_PENDING:
            try:
                with open(self.USER_MAP_FP, "w", encoding="utf8") as f:
                    json.dump(self.USER_MAP, f, ensure_ascii=False, indent=2)
                self.SAVE_PENDING = False
                logger.info(f"User map saved to {os.path.basename(self.USER_MAP_FP)}")
            except Exception as e:
                logger.error(f"Error saving user map file: {e}")
                QMessageBox.warning(self, "Warning", f"Could not save user dictionary: {e}")

    # --- Status update ---
    def update_status(self):
        # Check if status bar widgets exist before updating
        if hasattr(self, 'lineCol') and hasattr(self, 'wordCount') and hasattr(self, 'status'):
            cur = self.editor.textCursor()
            self.lineCol.setText(f"Ln {cur.blockNumber()+1}, Col {cur.columnNumber()+1}")
            words = len(WORD_PATTERN.findall(self.editor.toPlainText()))
            self.wordCount.setText(f"Words: {words}")

    # --- Show Event ---
    def showEvent(self, event):
        """Handle window show event to ensure it fits on screen."""
        super().showEvent(event)
        
        # Set a flag to indicate we're in the initial show event
        # This will prevent manual resizes from being overridden
        self._in_initial_show = True
        
        # Ensure window fits on screen after a short delay
        # to allow the window to fully initialize
        QTimer.singleShot(200, self.ensure_window_fits_screen)
        
        # Ensure keyboard fits on screen, but only during initial startup
        QTimer.singleShot(300, lambda: self._initial_keyboard_check())
        
        # Clear the initial show flag after a delay
        QTimer.singleShot(1000, lambda: setattr(self, '_in_initial_show', False))
        
    def _initial_keyboard_check(self):
        """Perform initial keyboard size check during startup only."""
        # Only run the keyboard size check if we're in the initial show
        if hasattr(self, '_in_initial_show') and self._in_initial_show:
            self.ensure_keyboard_fits_screen()
            
            # After initial sizing, set the keyboard as manually sized
            # to prevent automatic resizing from changing it
            self._manual_resize = True
            
            # Clear the manual resize flag after a delay
            QTimer.singleShot(1000, lambda: setattr(self, '_manual_resize', False))
    
    # --- Splitter Moved Event ---
    def on_splitter_moved(self, pos, index):
        """Handle splitter movement to update keyboard font size based on height."""
        # Get the current keyboard container height
        keyboard_height = self.keyboard_container.height()
        
        # Calculate a new font size based on the keyboard height using the helper method
        if hasattr(self.keyboard_area, 'height_for_font'):
            # Use the keyboard's helper method to calculate appropriate font size
            from ui.constants import BASE_KB_HEIGHT, BASE_KB_FONT, MIN_KB_FONT, MAX_KB_FONT
            
            # Calculate new font size proportional to height and round to nearest integer
            new_font_size = max(MIN_KB_FONT, min(MAX_KB_FONT, round(BASE_KB_FONT * keyboard_height / BASE_KB_HEIGHT)))
            
            # Update the preferences directly
            self.preferences["keyboard_font_size"] = new_font_size
            
            # Update the keyboard font size if needed
            if hasattr(self.keyboard_area, 'set_font_size'):
                self.keyboard_area.set_font_size(new_font_size)
                logging.info(f"Updated keyboard font size to {new_font_size} based on height {keyboard_height}")
        
        # Save the splitter state in preferences
        self.preferences["splitter_state"] = [pos, self.main_splitter.sizes()]
        
        # Save preferences to disk
        from app import config
        config.save_user_preferences(self.preferences)
    
    # --- Resize Event ---
    def resizeEvent(self, event):
        """Handle window resize events."""
        # Call the parent's resizeEvent
        super().resizeEvent(event)
        
        # Update suggestion popup position if it's visible
        if hasattr(self, 'suggestion_popup') and self.suggestion_popup.isVisible():
            self.update_suggestion_popup_position()
        
        # Update keyboard buttons to match the new window size
        if hasattr(self, 'keyboard_area'):
            try:
                # Update the keyboard buttons to match the new window size
                if hasattr(self.keyboard_area, 'update_buttons'):
                    self.keyboard_area.update_buttons()
            except Exception as e:
                logger.error(f"Error updating keyboard in resize event: {e}")
    
    # --- Close Event ---
    def closeEvent(self, event):
        """Handle application close event."""
        # Check for unsaved changes
        if self.editor.document().isModified():
            reply = QMessageBox.question(
                self, 
                "Save Changes?",
                "Do you want to save changes before exiting?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
            )

            if reply == QMessageBox.Cancel:
                event.ignore()
                return
            elif reply == QMessageBox.Yes:
                if not self.save_file():
                    # If save failed and user canceled, don't close
                    event.ignore()
                    return

        # Commit any pending buffer
        self.commit_buffer()

        # Save user dictionary
        self._save_map(force=True)

        # Save window size in preferences
        self.preferences["window_size"] = (self.width(), self.height())

        # Save font settings
        self.preferences["font"] = self.base_font.family()
        self.preferences["font_size"] = self.base_font.pointSize()

        # Save feature states
        self.preferences["show_keyboard"] = self.keyboard_area.isVisible()
        self.preferences["show_suggestions"] = self.suggestions_toggle_action.isChecked()
        self.preferences["singlish_enabled"] = self.singlish_toggle_action.isChecked()

        # Save all preferences
        config.save_user_preferences(self.preferences)
        logger.info("User preferences saved")

        # Remove event filter to prevent memory leaks
        self.editor.removeEventFilter(self)

        # Accept the close event
        event.accept()

    def print_document(self):
        """Handles printing the document."""
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            print_dialog = QPrintDialog(printer, self)
            print_dialog.setWindowTitle("Print Document")

            if print_dialog.exec() == QDialog.Accepted:
                document = self.editor.document()
                document.print_(printer)
                logger.info("Document sent to printer.")
                QMessageBox.information(self, "Print", "Document sent to printer.")
            else:
                logger.info("Print canceled by user.")
                QMessageBox.information(self, "Print", "Print canceled.")
        except Exception as e:
            logger.error(f"Error during printing: {e}")
            QMessageBox.critical(self, "Print Error", f"An error occurred while trying to print:\n{e}")

    def print_preview_document(self):
        """Shows a print preview dialog for the document."""
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            preview_dialog = QPrintPreviewDialog(printer, self)
            preview_dialog.setWindowTitle("Print Preview")
            preview_dialog.paintRequested.connect(self.handle_paint_request)
            
            # Set the size of the preview dialog (optional, but can be useful)
            # Get screen geometry to make it responsive
            screen_geometry = QApplication.primaryScreen().geometry()
            preview_dialog.resize(screen_geometry.width() * 0.7, screen_geometry.height() * 0.7)
            
            preview_dialog.exec()
            logger.info("Print preview dialog shown.")
        except Exception as e:
            logger.error(f"Error during print preview: {e}")
            QMessageBox.critical(self, "Print Preview Error", f"An error occurred while trying to show print preview:\n{e}")

    def handle_paint_request(self, printer):
        """Handles the paintRequested signal from QPrintPreviewDialog."""
        try:
            document = self.editor.document()
            document.print_(printer)
            logger.info("Document painted to print preview dialog.")
        except Exception as e:
            logger.error(f"Error painting document to preview: {e}")
            QMessageBox.critical(self, "Paint Error", f"An error occurred while painting the document for preview:\n{e}")

    def _load_dictionaries(self):
        """Loads the main lexicon and user-defined dictionary."""
        logger.info("Loading dictionaries...")
        try:
            # Load user map
            if os.path.exists(self.USER_MAP_FP):
                try:
                    with open(self.USER_MAP_FP, "r", encoding="utf8") as f:
                        self.USER_MAP = json.load(f)
                    logger.info(f"Loaded user map from {os.path.basename(self.USER_MAP_FP)}")
                except (IOError, json.JSONDecodeError) as e:
                    logger.error(f"Error loading user map: {e}")
                    QMessageBox.warning(self, "Warning", f"Could not load user dictionary: {e}")
                    self.USER_MAP = {}
            else:
                self.USER_MAP = {}
                logger.info(f"User map file not found at {os.path.basename(self.USER_MAP_FP)}. Starting with empty user map.")

            # Load main lexicon chunks
            self.MAIN_LEXICON = {}
            if os.path.exists(config.LEXICON_DIR):
                for filename in os.listdir(config.LEXICON_DIR):
                    if filename.endswith(".json.gz"):
                        filepath = os.path.join(config.LEXICON_DIR, filename)
                        try:
                            with gzip.open(filepath, "rt", encoding="utf8") as f:
                                chunk = json.load(f)
                                self.MAIN_LEXICON.update(chunk)
                            logger.info(f"Loaded lexicon chunk: {filename}")
                        except Exception as e:
                            logger.error(f"Error loading lexicon chunk {filename}: {e}")
            else:
                logger.warning(f"Lexicon chunks directory not found at {os.path.basename(config.LEXICON_DIR)}. Main lexicon will be empty.")

            # Combine user map and main lexicon for quick lookup
            self.MAP = {**self.MAIN_LEXICON, **self.USER_MAP}
            logger.info(f"Total entries loaded: {len(self.MAP)}")

        except Exception as e:
            logger.error(f"An error occurred during dictionary loading: {e}")
            QMessageBox.critical(self, "Error", f"Failed to load dictionaries: {e}")
    
    def show_settings_dialog(self):
        """Show the settings dialog."""
        dialog = SettingsDialog(self, self.preferences)
        dialog.settingsApplied.connect(self.apply_settings)
        if dialog.exec():
            logger.info("Settings applied")
    
    def apply_settings(self, settings):
        """Apply settings from the settings dialog."""
        # Update preferences
        for key, value in settings.items():
            self.preferences[key] = value
        
        # Save preferences
        config.save_user_preferences(self.preferences)
        
        # Apply font settings
        self.base_font = QFont(
            self.preferences["font"], 
            self.preferences["font_size"]
        )
        # Use PreferMatch to allow system fonts for English characters
        self.base_font.setStyleStrategy(QFont.StyleStrategy.PreferMatch)
        self.editor.setFont(self.base_font)
        
        # Apply keyboard settings - rewritten with consistent indentation
        if hasattr(self.keyboard_area, 'set_font_size'):
            # Use the proper method to update font size
            self.keyboard_area.set_font_size(self.preferences["keyboard_font_size"])
        else:
            # Fallback to direct property setting
            self.keyboard_area.font_size = self.preferences["keyboard_font_size"]
            self.keyboard_area.update_theme()  # This will update the button styles
        
        # Apply keyboard visibility
        if self.preferences["show_keyboard"]:
            self.keyboard_container.show()
            self.keyboard_area.show()
        else:
            self.keyboard_container.hide()
            self.keyboard_area.hide()
        
        # Update toggle actions
        self.singlish_toggle_action.setChecked(self.preferences["singlish_enabled"])
        self.singlish_toggle_action.setText("Singlish: On" if self.preferences["singlish_enabled"] else "Singlish: Off")
        
        self.suggestions_toggle_action.setChecked(self.preferences["show_suggestions"])
        self.suggestions_toggle_action.setText("Suggestions: On" if self.preferences["show_suggestions"] else "Suggestions: Off")
        
        self.keyboard_toggle_action.setChecked(self.preferences["show_keyboard"])
        self.keyboard_toggle_action.setText("Keyboard: On" if self.preferences["show_keyboard"] else "Keyboard: Off")
        
        logger.info("Settings applied successfully")


# ------------------------------------------------------------------
#  File Format Handlers
# ------------------------------------------------------------------
def read_text_file(file_path):
    """Read content from a plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def write_text_file(file_path, content):
    """Write content to a plain text file."""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

def read_docx_file(file_path):
    """Read content from a DOCX file."""
    try:
        import docx
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except ImportError:
        log.error("python-docx module not available")
        raise ImportError("Could not import python-docx module")

def write_docx_file(file_path, qtext_document):
    """Write content to a DOCX file with preserved formatting.
    
    Args:
        file_path: Path to save the DOCX file
        qtext_document: QTextDocument containing the formatted text
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        # Create a new document
        docx = Document()
        
        # If we received a string instead of a QTextDocument (for backward compatibility)
        if isinstance(qtext_document, str):
            # Fall back to the old behavior
            for paragraph in qtext_document.split('\n'):
                if paragraph.strip():  # Skip empty paragraphs
                    docx.add_paragraph(paragraph)
            docx.save(file_path)
            return
            
        # Process the QTextDocument to preserve formatting
        block = qtext_document.begin()
        
        while block != qtext_document.end():
            # Create a new paragraph for each text block
            p = docx.add_paragraph()
            
            # Get paragraph alignment
            block_format = block.blockFormat()
            if hasattr(block_format, 'alignment'):
                alignment = block_format.alignment()
                # Map Qt alignment to python-docx alignment
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if alignment == Qt.AlignLeft:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif alignment == Qt.AlignRight:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == Qt.AlignCenter:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == Qt.AlignJustify:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Process each text fragment in the block
            it = block.begin()
            while not it.atEnd():
                frag = it.fragment()
                if frag.isValid():
                    fmt = frag.charFormat()
                    text = frag.text()
                    
                    # Skip empty fragments
                    if text.strip():
                        run = p.add_run(text)
                        
                        # Set font family and size
                        if fmt.fontFamily():
                            run.font.name = fmt.fontFamily()
                        
                        # Set font size (default to 12pt if not specified)
                        font_size = fmt.fontPointSize()
                        if font_size > 0:
                            run.font.size = Pt(font_size)
                        else:
                            run.font.size = Pt(12)
                        
                        # Set bold, italic, underline
                        run.bold = fmt.fontWeight() > 50  # QFont.Normal is 50
                        run.italic = fmt.fontItalic()
                        run.underline = fmt.fontUnderline()
                        
                        # Set text color if valid
                        color = fmt.foreground().color()
                        if color.isValid():
                            run.font.color.rgb = RGBColor(color.red(), color.green(), color.blue())
                
                it += 1
            
            # Move to the next block
            block = block.next()
        
        # Save the document
        docx.save(file_path)
        
    except ImportError:
        log.error("python-docx module not available")
        raise ImportError("Could not import python-docx module")

def read_pdf_file(file_path):
    """Read content from a PDF file."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except ImportError:
        log.error("pypdf module not available")
        raise ImportError("Could not import pypdf module")

def write_pdf_file(file_path, qtext_document):
    """Write content to a PDF file using Qt's QPrinter."""
    try:
        from PySide6.QtPrintSupport import QPrinter
        from PySide6.QtGui import QPageSize # For standard page sizes

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)
       
        # Set page size to A4
        pageSize = QPageSize(QPageSize.PageSizeId.A4)
        printer.setPageSize(pageSize)
       
        # It's good practice to set page margins (optional, but often desired)
        # printer.setPageMargins(15, 15, 15, 15, QPrinter.Millimeter) # Example: 15mm margins

        qtext_document.print_(printer)
        log.info(f"Successfully wrote PDF to {file_path} using QPrinter")

    except ImportError:
        log.error("PySide6.QtPrintSupport QPrinter module not available")
        raise ImportError("Could not import QPrinter from PySide6.QtPrintSupport")
    except Exception as e:
        log.error(f"Error writing PDF with QPrinter: {e}")
        raise e # Re-raise the exception to be caught by the caller

# ------------------------------------------------------------------
#  Main function
# ------------------------------------------------------------------
def main():
    # Enable high DPI scaling
    QApplication.setHighDpiScaleFactorRoundingPolicy(Qt.HighDpiScaleFactorRoundingPolicy.PassThrough)

    # Create QApplication
    app = QApplication(sys.argv)
    
    # Set application icon
    icon_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                           "resources", "splash", "sinhalaword_icon.ico")
    if os.path.exists(icon_path):
        app.setWindowIcon(QIcon(icon_path))
    
    # Set application-wide font policy to allow system fonts for English text
    default_font = QFont()
    default_font.setStyleStrategy(QFont.StyleStrategy.PreferMatch)
    QApplication.setFont(default_font)

    # Create and show the main window
    win = SinhalaWordApp()

    # Set window size from preferences
    prefs = config.load_user_preferences()
    width, height = prefs["window_size"]
    win.resize(width, height)

    # Show the window
    win.show()

    # Start the application event loop
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
